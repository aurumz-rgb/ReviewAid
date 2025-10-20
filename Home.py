import streamlit as st
import fitz  
import time
import os 
import re
import json
import json5
import plotly.express as px
import pandas as pd
import cohere
from dotenv import load_dotenv
from datetime import datetime
from cryptography.fernet import Fernet
from fpdf import FPDF
from streamlit_lottie import st_lottie
import streamlit.components.v1 as components



st.set_page_config(
    page_title="ReviewAid",
    page_icon=os.path.abspath("favicon.ico"),
    layout="wide",
    initial_sidebar_state="expanded",
)





hide_streamlit_style = """
    <style>
    /* Hide hamburger menu */
    #MainMenu {visibility: hidden;}
    /* Hide footer */
    footer {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)





# Define your button colors (adjust as you want)
color = "#4189DC"  # blue
hover_color = "#174D8B"  # darker blue

# Put this CSS **after** your button is rendered (or just once at the top)
st.markdown(f"""
    <style>
    /* Target all Streamlit buttons */
    div.stButton > button:first-child {{
        background-color: {color} !important;
        color: white !important;
        border: none !important;
        font-weight: bold !important;
    }}
    div.stButton > button:first-child:hover {{
        background-color: {hover_color} !important;
    }}
    </style>
    """, unsafe_allow_html=True)






load_dotenv()

# Developer access controls
DEV_PASSWORD = os.getenv("DEV_PASSWORD", "defaultpassword")

if "dev_authenticated" not in st.session_state:
    st.session_state.dev_authenticated = False
if "show_dev_login" not in st.session_state:
    st.session_state.show_dev_login = False

with st.sidebar:
    if not st.session_state.dev_authenticated:
        if st.button("Developer Login"):
            st.session_state.show_dev_login = True

    if st.session_state.show_dev_login and not st.session_state.dev_authenticated:
        dev_pass = st.text_input("Enter developer password:", type="password")
        if st.button("Login as Developer"):
            if dev_pass == DEV_PASSWORD:
                st.session_state.dev_authenticated = True
                st.session_state.show_dev_login = False
                st.success("Developer access granted!")
                st.rerun()
            else:
                st.error("Incorrect password!")

    if st.session_state.dev_authenticated:
        st.markdown("### Developer Panel")

        unique_users = st.session_state.get("unique_users", set())
        user_api_keys = st.session_state.get("user_api_keys", set())
        papers_screened = st.session_state.get("papers_screened", 0)

        st.write(f"**Total unique users:** {len(unique_users)}")
        st.write(f"**Unique user API keys used:** {len(user_api_keys)}")
        st.write(f"**Total papers screened:** {papers_screened}")

        if st.button("Logout Developer"):
            st.session_state.dev_authenticated = False
            st.rerun()


ADMIN_API_KEY = os.getenv("ADMIN_API_KEY") or os.getenv("COHERE_API_KEY")

# --- Load Lottie ---
def load_lottiefile(filepath: str):
    try:
        with open(filepath, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return None

lottie_animation = load_lottiefile("animation.json")

# --- Session state ---
if "user_api_key" not in st.session_state:
    st.session_state.user_api_key = None
if "use_admin_key" not in st.session_state:
    st.session_state.use_admin_key = True
if "included_results" not in st.session_state:
    st.session_state.included_results = []
if "excluded_results" not in st.session_state:
    st.session_state.excluded_results = []
if "maybe_results" not in st.session_state:
    st.session_state.maybe_results = []
if "start_time" not in st.session_state:
    st.session_state.start_time = time.time()
if "note_acknowledged" not in st.session_state:
    st.session_state.note_acknowledged = False
if "disclaimer_acknowledged" not in st.session_state:
    st.session_state.disclaimer_acknowledged = False


# --- Helper ---
def init_cohere_client(api_key):
    return cohere.Client(api_key)

def test_api_key(api_key):
    if not api_key: return False
    try:
        client = cohere.Client(api_key)
        # Migrated to Chat API
        client.chat(model="command-r-08-2024", message="Hello", max_tokens=5)
        return True
    except Exception:
        return False

# ---------------- STAGE 1: Initial Note ----------------
if not st.session_state.note_acknowledged:
    st.markdown("""⚠️ **Note:**

- Upload limit is 20 papers maximum in a single session.  
- Do not re-screen the paper as it will consume 2x tokens.  
- This tool is in its **TRIAL** phase.  
- Confidence scores indicate text readability (max: 0.9).  

""")
    if st.button("Got it!"):
        st.session_state.note_acknowledged = True
        st.rerun()
    st.stop()  # Stop here if note not acknowledged

# ---------------- STAGE 2: Main App ----------------
# --- Show Lottie ABOVE heading ---
if lottie_animation:
    st.markdown("<div style='display:flex; justify-content:center;'>", unsafe_allow_html=True)
    st_lottie(lottie_animation, height=250, key="lottie_top")
    st.markdown("</div>", unsafe_allow_html=True)

# --- Heading ---

   
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@700&display=swap');

    .stApp {
        background-color: #0f1117;
        color: #F0F4F8;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    .typewriter {
        text-align: center;
        font-family: 'Montserrat', sans-serif;
        font-weight: 700;
        font-size: 3.5rem;
        color: #F0F4F8;
        margin-top: 10px;
    }

    .typewriter .typing {
        display: inline-block;
        overflow: hidden;
        border-right: .15em solid #F0F4F8;
        white-space: nowrap;
        letter-spacing: .04em;
        animation:
            typing 1s steps(11, end) forwards,
            blink-caret 1s step-end forwards;
        vertical-align: bottom;
    }

    .gold {
        color: #45a4f3;
    }

    .typewriter .dot {
        display: inline-block;
        vertical-align: bottom;
        margin-left: 2px;
        color: #F0F4F8;
    }

    @keyframes typing {
        from { width: 0 }
        to { width: 11ch; }
    }

    @keyframes blink-caret {
        0%, 100% { border-color: transparent; }
        50% { border-color: #F0F4F8; }
    }
    </style>

    <div class="typewriter">
        <span class="typing">Review<span class="gold">Aid</span>.</span>
    </div>

    <h3 style='text-align: center; color: #F0F4F8; margin-top: 10px;'>
        An Ai based Full-text Research Article Screener & Extractor
    </h3>
    """,
    unsafe_allow_html=True
)

# --- STEP 2: Check Admin Key ---
admin_key_valid = False
if st.session_state.use_admin_key:
    if test_api_key(ADMIN_API_KEY):
        admin_key_valid = True
    else:
        st.session_state.use_admin_key = False
        st.rerun()

# --- STEP 3: Disclaimer ---
if (st.session_state.user_api_key is None and not admin_key_valid) and not st.session_state.disclaimer_acknowledged:
    st.markdown("<div style='margin-top: 80px;'>", unsafe_allow_html=True)
    st.warning("""
    **Disclaimer:**  
    This tool is an independent educational resource and is **not affiliated with, endorsed by, or officially connected to Cohere**.  
    - You must use your **own Cohere API key** to access this tool.  
    - **Your API key is never stored, shared, or logged** by this tool.  
    - **You are solely responsible for securing and managing your API key**. Make sure to keep your API key safe and do not share it publicly.  
    - I do not take any responsibility for the use or misuse of the API key you obtain. Please use it responsibly and comply with Cohere’s policies.
    - By using this tool, you agree to comply with [Cohere's Terms of Service](https://cohere.com/terms-of-use) and [Acceptable Use Policy](https://docs.cohere.com/docs/cohere-labs-acceptable-use-policy).  
    """)
    agree = st.checkbox("I have read and agree to the above disclaimer.")
    if agree and st.button("I Agree & Continue"):
        st.session_state.disclaimer_acknowledged = True
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
    st.stop()

# --- STEP 4: Initialize Client ---
if st.session_state.user_api_key:
    co = init_cohere_client(st.session_state.user_api_key)
elif admin_key_valid:
    co = init_cohere_client(ADMIN_API_KEY)
else:
    st.markdown("<div style='height:80px;'></div>", unsafe_allow_html=True)
    st.markdown("""
    **Admin's Cohere API key has expired or is invalid.**  
    Please enter your own free Cohere API key below to continue. (https://cohere.com)  
    [How to get a free Cohere API key](./API_tutorial)
    """, unsafe_allow_html=True)
    user_key = st.text_input("Enter your free Cohere API key:", type="password")
    if user_key:
        if test_api_key(user_key):
            st.session_state.user_api_key = user_key
            st.success("Your API key is valid! You can continue.")
            st.rerun()
        else:
            st.error("Invalid API key, please try again.")
            st.stop()
    else:
        st.info("Please enter your API key to continue.")
        st.stop()



# FUNCTIONS 
def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    return "".join(page.get_text() for page in doc)

def preprocess_text_for_ai(text, max_tokens=1024):
    """Clean and truncate text."""
    text = " ".join(text.split())
    return text[:max_tokens * 4]  

def query_cohere(prompt):
    # Migrated to Chat API
    response = co.chat(
        model="command-r-08-2024",
        message=prompt,
        max_tokens=1024,
        temperature=0.2
    )
    return response.text




def extract_json_substring(text):
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end < start:
        return text  # no clear JSON braces found, return as is
    return text[start:end+1]

def repair_json_via_ai(broken_json_str):
    fix_prompt = f"""
The following JSON output is invalid or malformed. Please fix it and return ONLY valid JSON, no extra text:

{broken_json_str}
"""
    fixed_raw = query_cohere(fix_prompt)  # reuse your existing AI call
    return fixed_raw




def parse_result(raw_result):
    """Parse AI output with pre-cleanup, json5 advanced repair + retry + fallback."""
    try:
        # Pre-clean raw output to extract JSON substring
        cleaned = extract_json_substring(raw_result)

        # First, try to parse as standard JSON
        return json.loads(cleaned)
    except json.JSONDecodeError:
        print("⚠️ JSON parsing failed on cleaned string. Trying json5 parser...")

        try:
            # Try lenient json5 parser to handle more malformed JSON
            return json5.loads(cleaned)
        except Exception as e:
            print(f"⚠️ json5 parser failed: {e}. Attempting manual repair...")

            repaired = cleaned.strip()

            # Ensure it starts and ends with braces
            if not repaired.startswith("{"):
                repaired = "{" + repaired
            if not repaired.endswith("}"):
                repaired += "}"

            # Add quotes around unquoted keys (e.g., key: -> "key":)
            repaired = re.sub(r'(\s*)([a-zA-Z0-9_]+):', r'\1"\2":', repaired)

            try:
                return json.loads(repaired)
            except json.JSONDecodeError:
                print("⚠️ Manual repair failed. Trying AI fix retry...")

                # Retry fix via AI
                fixed_raw = repair_json_via_ai(raw_result)

                # Extract JSON substring again
                fixed_cleaned = extract_json_substring(fixed_raw)

                try:
                    return json.loads(fixed_cleaned)
                except json.JSONDecodeError:
                    print("⚠️ AI fix retry failed. Using fallback response.")
                    return {
                        "status": "Error",
                        "reason": "Invalid or unparseable JSON response after retries.",
                        "extracted": {}
                    }



    
def estimate_confidence(text):
    if not text or len(text.strip()) < 30:
        return 0.2
    if "randomized" in text.lower():
        return 0.9
    return 0.6

def df_from_results(results):
    rows = []
    for r in results:
        row = {
            "Filename": r.get("filename", ""),
            "Status": r.get("status", "").capitalize(),
            "Confidence": r.get("confidence", "")
        }
        row.update(r.get("extracted", {}))
        if r.get("status", "").lower() == "exclude":
            row["Reason for Exclusion"] = r.get("reason", "")
        rows.append(row)
    return pd.DataFrame(rows)


# Export helper functions
def to_docx(df):
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    doc = Document()
    doc.add_heading('Exported Papers', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].width = Inches(2)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            para = row_cells[i].paragraphs[0]
            para.add_run(str(val))
            tc = row_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(OxmlElement('w:vAlign'))
            tcPr[-1].set(qn('w:val'), 'top')
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def to_pdf(df):
    from fpdf import FPDF
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Arial", size=8)
    col_width = 280 / len(df.columns)
    for col in df.columns:
        pdf.multi_cell(col_width, 6, str(col), border=1, align='C')
    pdf.ln(6)
    for _, row in df.iterrows():
        for val in row:
            pdf.multi_cell(col_width, 6, str(val), border=1)
        pdf.ln(6)
    return pdf.output(dest='S').encode('latin1')

def to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def to_excel(df):
    import io
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    buffer.seek(0)
    return buffer.getvalue()




# Helper function for exclusion criteria matching 
def find_exclusion_matches(text, exclusion_lists):
    """
    Returns a list of exclusion criteria strings that are found in the given text.
    Simple substring matching (case-insensitive).
    """
    matches = []
    for criteria in exclusion_lists:
        criteria = criteria.strip()
        if criteria and criteria.lower() in text.lower():
            matches.append(criteria)
    return matches

#  UI 

st.subheader("Population Criteria")
population_inclusion = st.text_area("Population Inclusion Criteria", placeholder="e.g. Adults aged 18–65 with MS")
population_exclusion = st.text_area("Population Exclusion Criteria", placeholder="e.g. Patients with comorbid autoimmune diseases")

st.subheader("Intervention Criteria")
intervention_inclusion = st.text_area("Intervention Inclusion Criteria", placeholder="e.g. Natalizumab treatment ≥ 6 months")
intervention_exclusion = st.text_area("Intervention Exclusion Criteria", placeholder="e.g. Dose outside approved range")

st.subheader("Comparison Criteria")
comparison_inclusion = st.text_area("Comparison Inclusion Criteria", placeholder="e.g. Placebo or no treatment")
comparison_exclusion = st.text_area("Comparison Exclusion Criteria", placeholder="e.g. Active comparator like interferon beta")

st.subheader("Outcome Criteria (Optional)")
outcome_criteria = st.text_area("Outcome Criteria", placeholder="e.g. Annualized relapse rate, disability progression")

fields = st.text_input("Fields to Extract (comma-separated)", placeholder="e.g. Type of Study, Year, Population, Outcome")
fields_list = [f.strip() for f in fields.split(",") if f.strip()]
uploaded_pdfs = st.file_uploader("Upload PDF Files", accept_multiple_files=True)





# STREAMLIT BUTTON ACTION 
if st.button("Screen & Extract"):

    if "unique_users" not in st.session_state:
       st.session_state.unique_users = set()
    if "user_api_keys" not in st.session_state:
       st.session_state.user_api_keys = set()
    if "papers_screened" not in st.session_state:
       st.session_state.papers_screened = 0

    st.session_state.papers_screened += min(len(uploaded_pdfs), 20)   

# Track current user
    if st.session_state.user_api_key:
       st.session_state.unique_users.add(st.session_state.user_api_key)
       st.session_state.user_api_keys.add(st.session_state.user_api_key)
    elif st.session_state.use_admin_key:
       st.session_state.unique_users.add("admin")


    # validations
    if not fields.strip():
        st.warning("Please enter at least one field to be extracted!")
        st.stop()
    if not uploaded_pdfs:
        st.warning("Please upload at least one PDF file.")
        st.stop()
    if not any([
        population_inclusion.strip(), population_exclusion.strip(),
        intervention_inclusion.strip(), intervention_exclusion.strip(),
        comparison_inclusion.strip(), comparison_exclusion.strip(),
        outcome_criteria.strip()
    ]):
        st.warning("Please enter at least one inclusion or exclusion criterion.")
        st.stop()

    # Reset previous results
    st.session_state.included_results, st.session_state.excluded_results, st.session_state.maybe_results = [], [], []


    max_papers = 20
    total_pdfs = min(len(uploaded_pdfs), max_papers)
    progress_bar = st.progress(0)

    # Process uploaded PDFs
    for idx, pdf in enumerate(uploaded_pdfs[:max_papers], 1):
        st.info(f"Processing: {pdf.name}")
        try:
            pdf.seek(0)
            text = extract_text_from_pdf(pdf)

            if not text.strip():
                st.warning(f"PDF '{pdf.name}' appears empty or unreadable. Skipping.")
                progress_bar.progress(idx / total_pdfs)
                continue

            # Preprocess text 
            text = preprocess_text_for_ai(text, max_tokens=1024)
            confidence = estimate_confidence(text)

            # PRE-AI EXCLUSION CHECK 
            # Collect all exclusion criteria into a list
            all_exclusions = []
            for block in [population_exclusion, intervention_exclusion, comparison_exclusion]:
                if block.strip():
                    all_exclusions.extend([c.strip() for c in block.split(",") if c.strip()])

            # Find matches
            matches = find_exclusion_matches(text, all_exclusions)

            if len(matches) >= 1:
                # Auto-exclude without sending to AI
                exclusion_reason = (
                    f"Auto-excluded because {len(matches)} exclusion criteria matched: {', '.join(matches)}"
                )
                result = {
                    "filename": pdf.name,
                    "status": "Exclude",
                    "reason": exclusion_reason,
                    "confidence": confidence,
                    "extracted": {}
                }
                st.session_state.excluded_results.append(result)
                st.warning(f"Auto-excluded {pdf.name}: {len(matches)} exclusion criteria matched")
                progress_bar.progress(idx / total_pdfs)
                continue  # Skip AI step

            #AI-BASED CONFIGURATION
            prompt = f"""

**Population**
Inclusion: {population_inclusion}
Exclusion: {population_exclusion}

**Intervention**
Inclusion: {intervention_inclusion}
Exclusion: {intervention_exclusion}

**Comparison**
Inclusion: {comparison_inclusion}
Exclusion: {comparison_exclusion}

**Outcomes (if relevant)**: {outcome_criteria}

Fields to extract: {', '.join(fields_list)}

For each extracted field, please provide a detailed answer consisting of **1-2 complete sentences**, explaining the relevant information in context from the paper.
Also, provide a detailed reason for your classification.

Paper text:
\"\"\"
{text}
\"\"\"
For each field, provide 3-4 complete sentences with detailed, non-empty answers.
Give a detailed reason for classification ("Include", "Exclude", or "Maybe").

Return exactly one valid JSON object with this format, no extra text or comments:

{{
  "status": "Include",
  "reason": "Detailed classification reason.",
  "extracted": {{
    {', '.join(f'"{field}": "Detailed answer for {field}."' for field in fields_list)}
  }}
}}
"""
            with st.spinner(f"Sending '{pdf.name}' to AI..."):
                raw_result = query_cohere(prompt)

           

            result = parse_result(raw_result)
            if result.get("status") == "Error":
              st.warning(f"⚠️ Parsing failed for {pdf.name}. Using fallback response.")

            # Add metadata
            result["filename"] = pdf.name
            result["confidence"] = confidence
            if confidence < 0.5:
                result["flags"] = ["low_confidence"]

            # Fix: assign status variable for use below
            status = result.get("status", "").strip().lower()
            if status not in {"include", "exclude", "maybe"}:
                 status = "exclude"
           
            # Sort into appropriate result bucket
            if status == "include":
                st.session_state.included_results.append(result)
            elif status == "exclude":
                st.session_state.excluded_results.append(result)
            elif status == "maybe":
                st.session_state.maybe_results.append(result)
            else:
                st.session_state.excluded_results.append(result)

            st.success(f"Processed: {pdf.name} — {status.capitalize()}")

        except Exception as e:
            st.error(f"Error processing {pdf.name}: {str(e)}")

        progress_bar.progress(idx / total_pdfs)
        time.sleep(0.5)

    st.info("All files processed!")



included = len(st.session_state.included_results)
excluded = len(st.session_state.excluded_results)
maybe = len(st.session_state.maybe_results)
total = included + excluded + maybe

session_duration = time.time() - st.session_state.start_time
avg_speed = session_duration / total if total > 0 else 0




with st.expander("Screening Dashboard", expanded=True):
     st.metric("Papers Screened", total)

     fig = px.pie(
        names=["Included", "Excluded", "Maybe"],
        values=[included, excluded, maybe],
        title="Screening Decisions"
    )
     st.plotly_chart(fig, use_container_width=True)

                

    # Show summary tables
included_results = st.session_state.included_results
excluded_results = st.session_state.excluded_results
maybe_results = st.session_state.maybe_results 

if included_results:
    st.header("Included Papers")
    df_inc = df_from_results(included_results)
    st.dataframe(df_inc)
else:
    st.info("No Included papers found.")

if excluded_results:
    st.header("Excluded Papers")
    df_exc = df_from_results(excluded_results)
    st.dataframe(df_exc)
else:
    st.info("No Excluded papers found.")

    maybe_results = st.session_state.maybe_results

if maybe_results:
    st.header("Maybe Papers")
    df_maybe = df_from_results(maybe_results)
    st.dataframe(df_maybe)
else:
    st.info("No Maybe papers found.")

# Export options # Export options (DOCX, CSV, XLSX)
if included_results or excluded_results or maybe_results:
    st.header("Export Results")

    def export_buttons(df, label_prefix):
        formats = {
            "DOCX": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", to_docx(df), f"{label_prefix.lower()}_papers.docx"),
            "CSV":  ("text/csv", to_csv(df), f"{label_prefix.lower()}_papers.csv"),
            "XLSX": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", to_excel(df), f"{label_prefix.lower()}_papers.xlsx")
        }

        for fmt, (mime, data, filename) in formats.items():
            st.download_button(
                label=f"Download {label_prefix} as {fmt}",
                data=data,
                file_name=filename,
                mime=mime
            )

    if included_results:
        st.subheader("Included Papers")
        df_inc = df_from_results(included_results)
        export_buttons(df_inc, "Included")

    if excluded_results:
        st.subheader("Excluded Papers")
        df_exc = df_from_results(excluded_results)
        export_buttons(df_exc, "Excluded")

    if maybe_results:
        st.subheader("Maybe Papers")
        df_maybe = df_from_results(maybe_results)
        export_buttons(df_maybe, "Maybe")




st.markdown(
    """
    <style>
        .github-link {
            display: inline-block;
            color: #4189DC;
            
            font-size: 16px;
            text-decoration: underline;
            cursor: pointer;
            transition: color 0.3s ease;
        }
        .github-link:hover {
            color: #174D8B;
        }
        .github-container {
            text-align: right;
            margin-top: 100px;
            margin-bottom: 80px;
        }
    </style>
    <div class="github-container">
        <a href="https://github.com/aurumz-rgb/ReviewAid" target="_blank" class="github-link">
             GitHub Source Code
        </a>
    </div>
    """,
    unsafe_allow_html=True
)





import os
from datetime import datetime


version = "1.0.2"
try:
    current_file = __file__
    last_modified_timestamp = os.path.getmtime(current_file)
    last_updated = datetime.fromtimestamp(last_modified_timestamp).strftime("%Y-%m-%d %H:%M:%S")
except Exception:
    last_updated = "Unknown"

# Footer
st.markdown(
    f"""
    <style>
    .custom-footer-container {{
        width: 100%;
        font-family: 'Times New Roman', Times, serif;
        font-size: 11px;  /* thinner font */
        color: #F0F4F8;
        opacity: 0.8;
        padding: 3px 6px;  /* less padding for thinner footer */
        position: fixed;
        bottom: 0;
        left: 0;
        background-color: #1f2937;
        display: flex;
        justify-content: space-between;  /* spread left & right */
        align-items: center;
        z-index: 9999;
    }}
    </style>

    <div class="custom-footer-container">
        <div style="white-space: nowrap; letter-spacing: 1px;">
            Version {version}
        </div>
        <div style="white-space: nowrap; letter-spacing: 1px;">
            Made with 💙 by its Creator.
        </div>
    </div>
    """,
    unsafe_allow_html=True
)
import streamlit as st
import fitz  
import time
import os 
import re
import json
import json5
import plotly.express as px
import pandas as pd
from dotenv import load_dotenv
from datetime import datetime
from cryptography.fernet import Fernet
from fpdf import FPDF
from streamlit_lottie import st_lottie
import streamlit.components.v1 as components
import base64
import html
from zai import ZaiClient
import gc 
import uuid 
import hashlib
import firebase_admin
from firebase_admin import credentials, firestore
import io  

load_dotenv()


db = None 

try:
    if not firebase_admin._apps:
        firebase_key_str = os.getenv("FIREBASE_KEY")
        if firebase_key_str:
           
            cred_dict = json.loads(firebase_key_str)
            cred = credentials.Certificate(cred_dict)
            firebase_admin.initialize_app(cred)
    

    if firebase_admin._apps:
        db = firestore.client()

except Exception as e:
    st.error(f"Error initializing Firebase: {e}")

def increment_firebase_counter(field):
    """Safely increments a counter in Firestore."""
    if db is None:
        return
    try:
        doc_ref = db.collection("ReviewAidAnalytics").document("counters")
        doc_ref.update({field: firestore.Increment(1)})
    except Exception as e:
        print(f"Analytics Error: {e}")

def get_firebase_stats():
    """Reads counters from Firestore."""
    if db is None:
        return {
            "papers_screened": 0,
            "papers_extracted": 0,
            "total_visits": 0,
        }
    try:
        doc_ref = db.collection("ReviewAidAnalytics").document("counters")
        doc = doc_ref.get()
        if doc.exists:
            return doc.to_dict()
        else:
            return {
                "papers_screened": 0,
                "papers_extracted": 0,
                "total_visits": 0,
            }
    except Exception as e:
        print(f"Analytics Read Error: {e}")
        return {
            "papers_screened": 0,
            "papers_extracted": 0,
            "total_visits": 0,
        }

def init_analytics():
    """Initializes session and updates global counters via Firebase."""
    if db is None:
        return
    if "visit_recorded" not in st.session_state:
        increment_firebase_counter("total_visits")
        st.session_state.visit_recorded = True

def update_processing_stats(mode, count=1):
    """Updates papers screened/extracted count via Firebase."""
    if db is None:
        return
    if mode == "screener":
        for _ in range(count):
            increment_firebase_counter("papers_screened")
    elif mode == "extractor":
        for _ in range(count):
            increment_firebase_counter("papers_extracted")

st.set_page_config(
    page_title="ReviewAid / AI Screener & Extractor",
    page_icon=os.path.abspath("favicon.ico"),
    layout="wide",
    initial_sidebar_state="collapsed",
)

hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    [data-testid="stSidebar"] {display: none;}
    
    .main .block-container {
        padding-bottom: 0 !important;
        margin-bottom: 0 !important;
        padding-top: 0 !important;
        margin-top: -30px !important;
    }
    
    div[data-testid="stVerticalBlock"] > div:empty {
        display: none !important;
    }
    
    div[data-testid="stVerticalBlock"]:empty {
        display: none !important;
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    
    .citation-box {
        background-color: rgba(65, 137, 220, 0.1);
        border-left: 4px solid #4189DC;
        padding: 15px;
        border-radius: 5px;
        margin: 15px 0;
    }
    
    .citation-box p {
        white-space: pre-line;
    }
    
    .support-links {
        display: flex;
        gap: 30px;
        margin: 20px 0;
    }
    
    .support-link-item {
        display: flex;
        flex-direction: column;
    }
    
    .support-link-item a {
        color: #4189DC;
        text-decoration: none;
        font-weight: bold;
        font-size: 1.1rem;
        transition: color 0.3s ease;
    }
    
    .support-link-item a:hover {
        color: #174D8B;
    }
    
    .support-link-item p {
        color: #d1d5db;
        margin-top: 5px;
        font-size: 0.9rem;
    }
    
    .custom-button {
        background-color: #4189DC !important;
        color: white !important;
        text-decoration: none !important;
        padding: 0.5rem 1rem !important;
        font-size: 1rem !important;
        border-radius: 5px !important;
        display: inline-block !important;
        margin: 0 !important;
        transition: background-color 0.3s ease !important;
        border: none !important;
        cursor: pointer !important;
    }
    
    .custom-button:hover {
        background-color: #174D8B !important;
    }
    
    .bold-white-link {
        color: #F0F4F8 !important;
        font-weight: bold !important;
        text-decoration: none !important;
        transition: color 0.3s ease !important;
    }
    
    .bold-white-link:hover {
        color: #45a4f3 !important;
    }
    
    .support-section {
        margin-top: 40px;
        margin-bottom: 30px;
    }
    
    .support-description {
        color: #d1d5db;
        margin-bottom: 20px;
        font-size: 1.2rem;
        line-height: 1.6;
    }
    
    .support-description a {
        color: #4189DC !important;
        text-decoration: none;
        font-weight: bold;
        transition: color 0.3s ease;
    }
    
    .support-description a:hover {
        color: #174D8B !important;
        text-decoration: underline;
    }
    
    .disclaimer-warning {
        background-color: #ccd0d9;
        color: #000000;
        border-left: 5px solid #174D8B;
        padding: 15px;
        border-radius: 5px;
        margin: 15px 0;
    }
    
    .disclaimer-warning h3 {
        margin-top: 0;
        color: #000000;
    }
    
    .disclaimer-warning p {
        margin-bottom: 10px;
        color: #000000;
    }
    
    .disclaimer-warning ul {
        margin-top: 10px;
        margin-bottom: 10px;
        padding-left: 20px;
    }
    
    .disclaimer-warning li {
        margin-bottom: 8px;
        color: #000000;
    }

    .terminal-container {
        background-color: #0d1117; 
        color: #c9d1d9;
        font-family: 'Courier New', Courier, monospace;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #30363d;
        height: 400px; 
        overflow-y: auto;
        font-size: 13px;
        line-height: 1.4;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.6);
        scroll-behavior: smooth;
    }
    
    .terminal-line {
        margin-bottom: 2px;
        padding-bottom: 1px;
        display: block;
        white-space: pre-wrap;
    }
    
    .terminal-timestamp {
        color: #8b949e;
        margin-right: 10px;
        font-weight: bold;
        min-width: 70px;
        display: inline-block;
    }
    
    .log-info { color: #58a6ff; }
    .log-success { color: #3fb950; }
    .log-warn { color: #d29922; }
    .log-error { color: #f85149; }
    .log-system { color: #a371f7; font-weight: bold;}
    .log-debug { color: #6e7681; font-style: italic; font-size: 12px; }
    
    .terminal-container::-webkit-scrollbar {
        width: 10px;
    }
    .terminal-container::-webkit-scrollbar-track {
        background: #0d1117;
    }
    .terminal-container::-webkit-scrollbar-thumb {
        background: #30363d;
        border-radius: 5px;
    }
    .terminal-container::-webkit-scrollbar-thumb:hover {
        background: #484f58;
    }

    .important-note-box {
        background-color: rgba(65, 137, 220, 0.15);
        border-left: 5px solid #e64d43;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 20px;
        color: #F0F4F8;
        font-size: 1rem;
        line-height: 1.5;
    }

    .confidence-table-container {
        background-color: rgba(31, 41, 55, 0.2);
        border-radius: 8px;
        padding: 20px;
        margin-bottom: 30px;
        color: #F0F4F8;
    }

    .confidence-table-container th {
        text-align: center;
        color: #45a4f3;
        border-bottom: 2px solid #45a4f3;
    }
    
    .confidence-table-container td {
        vertical-align: top;
        text-align: center;
        padding: 8px 0;
        border-bottom: 1px solid rgba(65, 137, 220, 0.2);
    }

    @media screen and (max-width: 768px) {
        .main .block-container {
            margin-top: -10px !important;
            padding-top: 10px !important;
        }
        .terminal-container {
            height: 250px !important;
            font-size: 11px !important;
            padding: 10px !important;
        }
        .terminal-timestamp {
            min-width: 55px;
            font-size: 10px;
        }
        .disclaimer-warning {
            padding: 10px !important;
            font-size: 0.9rem !important;
        }
        .important-note-box {
            padding: 10px !important;
            font-size: 0.9rem !important;
        }
        .stTextArea, .stTextInput {
            font-size: 0.9rem !important;
        }
        .confidence-table-container {
            padding: 10px !important;
            overflow-x: auto;
        }
    }
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

color = "#4189DC"
hover_color = "#174D8B"

st.markdown(f"""
    <style>
    div.stButton > button:first-child {{
        background-color: {color} !important;
        color: white !important;
        border: none !important;
        font-weight: bold !important;
    }}
    div.stButton > button:first-child:hover {{
        background-color: {hover_color} !important;
    }}
    
    .mode-card {{
        background-color: #1f2937;
        border-radius: 10px;
        padding: 30px;
        margin: 15px;
        width: 100%;
        max-width: 500px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }}
    
    .mode-card:hover {{
        transform: translateY(-5px);
        box-shadow: 0 10px 20px rgba(0, 0, 0, 0.2);
    }}
    
    .mode-title {{
        font-size: 1.5rem;
        font-weight: bold;
        margin-bottom: 15px;
        color: #F0F4F8;
    }}
    
    .mode-description {{
        color: #d1d5db;
        margin-bottom: 20px;
    }}
    
    .mode-selection-footer {{
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        background-color: rgba(8, 25, 45, 0.95);
        color: #d1d5db;
        font-size: 0.95rem;
        padding: 10px 20px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        z-index: 1000;
        backdrop-filter: blur(5px);
    }}
    
    .mode-selection-footer a {{
        color: #F0F4F8;
        text-decoration: none;
        font-weight: normal;
    }}
    
    .mode-selection-footer a:hover {{
        text-decoration: underline;
    }}
    
    .mode-selection-footer .center-text {{
        font-weight: normal;
    }}
    
    .mode-selection-footer .right-text {{
        font-weight: normal;
    }}
    
    @media screen and (max-width: 768px) {{
        .mode-card {{
            padding: 15px !important;
            margin: 5px 0 !important;
            max-width: 100% !important;
        }}
        .mode-title {{
            font-size: 1.2rem !important;
        }}
        .mode-description {{
            font-size: 0.9rem !important;
        }}
        .mode-selection-footer {{
            flex-direction: column !important;
            text-align: center;
            gap: 5px;
            padding: 10px !important;
        }}
        .custom-button {{
            width: 100% !important;
            text-align: center !important;
            box-sizing: border-box;
        }}
    }}
    </style>
    """, unsafe_allow_html=True)

@st.cache_resource
def load_lottiefile(filepath: str):
    try:
        with open(filepath, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return None

lottie_animation = load_lottiefile("animation.json")


if "included_results" not in st.session_state:
    st.session_state.included_results = []
if "excluded_results" not in st.session_state:
    st.session_state.excluded_results = []
if "maybe_results" not in st.session_state:
    st.session_state.maybe_results = []
if "extracted_results" not in st.session_state:
    st.session_state.extracted_results = []
if "start_time" not in st.session_state:
    st.session_state.start_time = time.time()
if "disclaimer_acknowledged" not in st.session_state:
    st.session_state.disclaimer_acknowledged = False
if "app_mode" not in st.session_state:
    st.session_state.app_mode = None
if "page_load_count" not in st.session_state:
    st.session_state.page_load_count = 0
if "terminal_logs" not in st.session_state:
    st.session_state.terminal_logs = []
if "batch_file_hashes" not in st.session_state:
    st.session_state.batch_file_hashes = {}

MAX_LOG_ENTRIES = 200
MAX_INPUT_TOKENS_SCREENER = 128000 
MAX_INPUT_TOKENS_EXTRACTOR = 128000 
MAX_OUTPUT_TOKENS = 4096

st.session_state.page_load_count += 1

img_src = ""
try:
    img_path = os.path.join("assets", "RA.png")
    if os.path.exists(img_path):
        with open(img_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
        img_src = f"data:image/png;base64,{encoded_string}"
        del encoded_string, image_file 
    else:
        img_path_root = "RA.png"
        if os.path.exists(img_path_root):
             with open(img_path_root, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode()
             img_src = f"data:image/png;base64,{encoded_string}"
             del encoded_string, image_file 
except Exception:
    pass 

if img_src:
    st.markdown(
        f"""
        <img src="{img_src}" style="position: fixed; bottom: 38px; right: 5px; width: 80px; height: 80px; z-index: 10000; opacity: 0.6; pointer-events: none;">
        """, unsafe_allow_html=True
    )

def query_zai(prompt, api_key, temperature=0.1, max_tokens=2048):
    if not api_key:
        st.error("API key is missing. Please check your environment variables.")
        return None
    
    update_terminal_log(f"Initializing ZAI Client...", "DEBUG")
    
    
    max_retries = 10 
    
    for attempt in range(max_retries):
        try:
            update_terminal_log(f"API Call Attempt {attempt + 1}/{max_retries}...", "DEBUG")
            
            client = ZaiClient(api_key=api_key)
            
            response = client.chat.completions.create(
                model="GLM-4.6V-Flash",
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            result_content = response.choices[0].message.content

            del response 
            del client 
            
            update_terminal_log("Response received successfully.", "SUCCESS")
            return result_content
            
        except Exception as e:
            error_str = str(e)
            is_rate_limit = False
            if "429" in error_str or "rate limit" in error_str.lower() or "too many requests" in error_str.lower() or "quota" in error_str.lower():
                is_rate_limit = True
            
            if is_rate_limit:
                update_terminal_log(f"Rate Limit / Quota Exceeded detected.", "WARN")
                
            
                wait_time = 15 * (2 ** attempt)
                
                if attempt < max_retries - 1:
                    update_terminal_log(f"Rate limit hit. Waiting {wait_time} seconds before retry...", "WARN")
                    time.sleep(wait_time)
                    update_terminal_log(f"Resuming retry...", "INFO")
                else:
                    update_terminal_log("Max retries reached for rate limit. Falling back to local processing.", "ERROR")
                    return None
            else:
                update_terminal_log(f"API Error: {error_str}", "ERROR")
             
                if attempt < 3:
                     time.sleep(2)
                else:
                     return None
    
    return None

def clean_json_response(raw_str):
    """
    Bulletproof JSON cleaning pipeline.
    Handles Markdown, Trailing Commas, Comments, and Control Characters.
    """
    if not raw_str:
        return ""

    raw_str = re.sub(r'```json\s*', '', raw_str)
    raw_str = re.sub(r'```\s*', '', raw_str)
    
    raw_str = re.sub(r'//.*', '', raw_str)
    
    raw_str = re.sub(r'/\*.*?\*/', '', raw_str, flags=re.DOTALL)
    
    raw_str = re.sub(r',\s*([}\]])', r'\1', raw_str)
    
    start = raw_str.find('{')
    end = raw_str.rfind('}')
    
    if start == -1 or end == -1 or end < start:
        return "" 
    
    cleaned = raw_str[start:end+1]
    
    def replace_newlines_in_strings(match):
        return match.group(0).replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')

    cleaned = re.sub(r'"(?:\\.|[^"\\])*"', replace_newlines_in_strings, cleaned)
    
    return cleaned

def parse_result(raw_result, api_key, mode="screener", fields_list=None, original_text=None):
    """
    Parses AI response with extreme prejudice.
    Tries Standard JSON -> JSON5 -> AI Repair -> Re-extraction -> Regex Fallback.
    """
    if raw_result is None:
        update_terminal_log("API response was None. Using Regex Fallback (Local Processing).", "WARN")
        if original_text:
            return _regex_extract_fallback(original_text, mode, fields_list)
        else:
            return _get_default_result(mode, fields_list)
    
    update_terminal_log("Starting Bulletproof JSON parsing pipeline...", "DEBUG")
    
    cleaned_json = clean_json_response(raw_result)
    
    if not cleaned_json or len(cleaned_json) < 10:
        update_terminal_log("Cleaned JSON is empty or invalid. Structure likely missing.", "WARN")
        if original_text:
            update_terminal_log("Attempting Re-extraction due to structural failure...", "WARN")
            res = _attempt_re_extraction(original_text, api_key, mode, fields_list)
            del original_text 
            return res
        return _regex_extract_fallback(raw_result, mode, fields_list)

    try:
        update_terminal_log("Attempting standard json.loads()...", "DEBUG")
        data = json.loads(cleaned_json)
        update_terminal_log("Standard JSON parse successful.", "SUCCESS")
        del cleaned_json 
        return data
    except json.JSONDecodeError as e:
        update_terminal_log(f"Standard JSON failed: {str(e)}", "WARN")

    try:
        update_terminal_log("Attempting JSON5 parser (relaxed standard)...", "DEBUG")
        data = json5.loads(cleaned_json)
        update_terminal_log("JSON5 parse successful.", "SUCCESS")
        del cleaned_json 
        return data
    except Exception as e:
        update_terminal_log(f"JSON5 failed: {str(e)}", "WARN")

    update_terminal_log("Attempting AI-based JSON repair...", "WARN")
    repair_prompt = f"""
The system generated a malformed JSON response. Your task is to fix the syntax errors and return ONLY a valid JSON object.

Rules:
1. Do NOT change the values, only fix the syntax (quotes, commas, braces).
2. Do NOT include markdown formatting (```).
3. Return ONLY the JSON.

Malformed JSON:
{cleaned_json}
"""
    fixed_raw = query_zai(repair_prompt, api_key, temperature=0.1, max_tokens=1024)
    del repair_prompt 
    
    if fixed_raw and fixed_raw != "RATE_LIMIT_ERROR":
        fixed_cleaned = clean_json_response(fixed_raw)
        if fixed_cleaned:
            try:
                update_terminal_log("Testing AI-Repaired JSON...", "DEBUG")
                data = json.loads(fixed_cleaned)
                update_terminal_log("AI repair successful.", "SUCCESS")
                del fixed_raw, fixed_cleaned
                return data
            except:
                update_terminal_log("AI repair failed.", "ERROR")
    else:
        update_terminal_log("AI repair skipped (empty/rate limit).", "WARN")

    update_terminal_log("All parsers failed. Using Regex Extraction Fallback.", "ERROR")
    return _regex_extract_fallback(raw_result, mode, fields_list)

def _attempt_re_extraction(original_text, api_key, mode, fields_list):
    """
    If first attempt failed (empty or bad structure), try once more with a very strict prompt.
    """
    update_terminal_log("Re-extraction initiated...", "SYSTEM")
    

    text_snippet = original_text[:MAX_INPUT_TOKENS_SCREENER*3] 
    
    if mode == "screener":
        strict_prompt = f"""
You failed to provide a valid JSON response previously. You are an expert systematic reviewer.
Analyze this text and return a valid JSON object ONLY.

Text:
\"\"\"{text_snippet}\"\"\"

Required JSON Format:
{{
  "status": "Include/Exclude/Maybe",
  "reason": "Brief reason",
  "title": "Paper Title",
  "author": "Author Name",
  "year": "Year",
  "confidence": 0.5
}}
Return ONLY JSON object.
"""
    else:
        fields_str = ", ".join(fields_list)
        strict_prompt = f"""
You failed to provide a valid JSON response previously.
Extract these fields: {fields_str}.
If a field is not found, use the value "Not Found".

Text:
\"\"\"{text_snippet}\"\"\"

Required JSON Format:
{{
  "extracted": {{
    "{fields_list[0]}": "Value",
    ...
  }},
  "confidence": 0.5
}}
Return ONLY JSON object.
"""

    re_raw = query_zai(strict_prompt, api_key, temperature=0.1, max_tokens=2048)
    del strict_prompt
    del text_snippet 
    
    if re_raw and re_raw != "RATE_LIMIT_ERROR":
        cleaned = clean_json_response(re_raw)
        try:
            data = json.loads(cleaned)
            update_terminal_log("Re-extraction successful.", "SUCCESS")
            del re_raw, cleaned
            return data
        except:
            pass
    
    update_terminal_log("Re-extraction failed. Using default/regex.", "ERROR")
    return _get_default_result(mode, fields_list)

def _get_default_result(mode, fields_list):
    if mode == "screener":
        return {
            "status": "Error",
            "reason": "Failed to extract data",
            "title": "Not Found",
            "author": "Not Found",
            "year": "Not Found",
            "confidence": 0.0
        }
    else:
        result = {"extracted": {}, "confidence": 0.0}
        if fields_list:
            for field in fields_list:
                result["extracted"][field] = "Not Found"
        return result

def _regex_extract_fallback(text, mode, fields_list):
    """
    Extracts specific key-value pairs from unstructured text using Regex.
    Used when JSON parsing fails completely. Improved to handle non-JSON formats.
    """
    update_terminal_log("Running Regex Fallback extraction...", "DEBUG")
    
    result = {}
    confidence = 0.2 
    
    if mode == "screener":
        result["status"] = "Error"
        result["reason"] = "Regex Fallback: AI Blocked - Using Local Rules"
        result["title"] = "Not Found"
        result["author"] = "Not Found"
        result["year"] = "Not Found"
        confidence = 0.2 

        lower_t = text.lower()
        
        if "include" in lower_t and "exclude" not in lower_t:
            result["status"] = "Include"
            result["reason"] = "Regex Fallback: Inferred Inclusion (Local)"
            confidence = 0.3
        elif "exclude" in lower_t:
            result["status"] = "Exclude"
            result["reason"] = "Regex Fallback: Inferred Exclusion (Local)"
            confidence = 0.3
        
        patterns = {
            "title": [r'"title"\s*:\s*"([^"]+)"', r'title\s*:\s*"?([^"\n]+)"?', r'Title\s*[:\-]\s*([^\n]+)'],
            "author": [r'"author"\s*:\s*"([^"]+)"', r'author\s*:\s*"?([^"\n]+)"?'],
            "year": [r'"year"\s*:\s*"([^"]+)"', r'year\s*:\s*(\d{4})']
        }
        
        for key, regex_list in patterns.items():
            for pattern in regex_list:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    val = match.group(1).strip()
                    result[key] = val
                    break 
        
    else:
        result = {"extracted": {}, "confidence": 0.2} 
        if fields_list:
            for field in fields_list:
                val = "Not Found"
                pattern = rf'"{field}"\s*:\s*"([^"]*)"'
                match = re.search(pattern, text, re.IGNORECASE)
                if not match:
                    pattern = rf'{field}\s*[:\-]\s*"?([^"\n]+)"?'
                    match = re.search(pattern, text, re.IGNORECASE)
                
                if match:
                    val = match.group(1).strip()
                
                result["extracted"][field] = val
        result["confidence"] = confidence
    
    return result

def display_citation_section():
    st.markdown("---")
    st.markdown("## Citation")

    apa_citation = (
        "Sahu, V. (2025). ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0). "
        "Zenodo. https://doi.org/10.5281/zenodo.18060973"
    )

    harvard_citation = (
        "Sahu, V., 2025. ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0). "
        "Zenodo. Available at: https://doi.org/10.5281/zenodo.18060973"
    )

    mla_citation = (
        "Sahu, Vihaan. \"ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0).\" "
        "2025, Zenodo, https://doi.org/10.5281/zenodo.18060973."
    )

    chicago_citation = (
        "Sahu, Vihaan. 2025. \"ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0).\" "
        "Zenodo. https://doi.org/10.5281/zenodo.18060973."
    )

    ieee_citation = (
        "V. Sahu, \"ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0),\" "
        "Zenodo, 2025. doi: 10.5281/zenodo.18060973."
    )

    vancouver_citation = (
        "Sahu V. ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0). "
        "Zenodo. 2025. doi:10.5281/zenodo.18060973"
    )

    ris_data = """TY  - JOUR
AU  - Sahu, V
TI  - ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0)
PY  - 2025
DO  - 10.5281/zenodo.18060973
ER  -"""

    bib_data = """@misc{Sahu2025,
  author={Sahu, V.},
  title={ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.0.0)},
  year={2025},
  doi={10.5281/zenodo.18060973}
}"""

    citation_style = st.selectbox(
        "Select citation style",
        ["APA", "Harvard", "MLA", "Chicago", "IEEE", "Vancouver"]
    )

    if citation_style == "APA":
        citation_text = apa_citation
    elif citation_style == "Harvard":
        citation_text = harvard_citation
    elif citation_style == "MLA":
        citation_text = mla_citation
    elif citation_style == "Chicago":
        citation_text = chicago_citation
    elif citation_style == "IEEE":
        citation_text = ieee_citation
    elif citation_style == "Vancouver":
        citation_text = vancouver_citation

    escaped_citation = html.escape(citation_text)
    
    st.markdown(f'<div class="citation-box"><p style="margin:0; color: #F0F4F8;">{escaped_citation}</p></div>', unsafe_allow_html=True)

    js_citation_text = json.dumps(citation_text)
    
    st.markdown(f"""
    <div style="display:flex; gap:10px; margin-top:10px; margin-bottom:10px; position:relative; flex-wrap:wrap;" id="button-container">
        <button id="copy-btn" class="custom-button">Copy</button>
        <a download="ReviewAid_citation.ris" href="data:application/x-research-info-systems;base64,{base64.b64encode(ris_data.encode()).decode()}" class="custom-button">RIS Format</a>
        <a download="ReviewAid_citation.bib" href="data:application/x-bibtex;base64,{base64.b64encode(bib_data.encode()).decode()}" class="custom-button">BibTeX Format</a>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <script>
    function copyCitation() {{
        const citationText = {js_citation_text};
        navigator.clipboard.writeText(citationText).then(() => {{
            const btn = document.getElementById('copy-btn');
            const originalText = btn.innerText;
            btn.innerText = "Copied!";
            setTimeout(() => {{ 
                btn.innerText = originalText; 
            }}, 2000);
        }}).catch(err => {{
            console.error('Failed to copy text: ', err);
        }});
    }}
    
    document.addEventListener('DOMContentLoaded', function() {{
        const copyBtn = document.getElementById('copy-btn');
        if (copyBtn) {{
            copyBtn.addEventListener('click', copyCitation);
        }}
    }});
    </script>
    """, unsafe_allow_html=True)

init_analytics()

if st.session_state.app_mode is not None:
    st.markdown("""
    <div class="return-button">
    """, unsafe_allow_html=True)
    if st.button("← Return to Mode Selection"):
        st.session_state.app_mode = None
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

st.markdown("<div class='content-wrapper'>", unsafe_allow_html=True)

if lottie_animation:
    st.markdown("<div class='lottie-container' style='display:flex; justify-content:center;'>", unsafe_allow_html=True)
    st_lottie(lottie_animation, height=250, key=f"lottie_top_{st.session_state.page_load_count}")
    st.markdown("</div>", unsafe_allow_html=True)

st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@700&display=swap');

    .stApp {
        background-color: #0f1117;
        color: #F0F4F8;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    .typewriter {
        text-align: center;
        font-family: 'Montserrat', sans-serif;
        font-weight: 700;
        font-size: 3.5rem;
        color: #F0F4F8;
    }

    .typewriter .typing {
        display: inline-block;
        overflow: hidden;
        border-right: .15em solid #F0F4F8;
        white-space: nowrap;
        letter-spacing: .04em;
        animation:
            typing 1s steps(11, end) forwards,
            blink-caret 1s step-end forwards;
        vertical-align: bottom;
    }

    .gold {
        color: #45a4f3;
    }

    .typewriter .dot {
        display: inline-block;
        vertical-align: bottom;
        margin-left: 2px;
        color: #F0F4F8;
    }

    @keyframes typing {
        from { width: 0 }
        to { width: 11ch; }
    }

    @keyframes blink-caret {
        0%, 100% { border-color: transparent; }
        50% { border-color: #F0F4F8; }
    }

    @media screen and (max-width: 768px) {
        .typewriter {
            font-size: 2.2rem !important;
        }
        .subheading {
            font-size: 1rem !important;
            margin-bottom: 10px;
        }
        .lottie-container svg {
            height: 180px !important;
        }
    }
    </style>

    <div class="typewriter">
        <span class="typing">Review<span class="gold">Aid</span>.</span>
    </div>

    <h3 class="subheading" style='text-align: center; color: #F0F4F8; margin-top: 0; margin-bottom: 5px;'>
        An Ai based Full-text Research Article Screener & Extractor
    </h3>
    """,
    unsafe_allow_html=True
)

if not st.session_state.disclaimer_acknowledged:
    st.markdown("<div style='margin-top: 80px;'>", unsafe_allow_html=True)
    
    with st.expander("Read Full Disclaimer"):
        st.markdown("""
        <div class="disclaimer-warning">
            <h3>ReviewAid Disclaimer:</h3>
    <p>By using ReviewAid, you acknowledge and agree to following:</p>
    <ul>
        <li><strong>Researcher Responsibility:</strong> ReviewAid is an AI-assisted tool intended to support, not replace, researcher's own judgment. All screening decisions, extracted data, and interpretations generated by system must be independently verified by user. The developer is not responsible for inaccuracies, omissions, or misclassifications arising from AI-generated outputs.</li>
        <li><strong>No Guarantee of Completeness or Accuracy:</strong> While ReviewAid aims to improve efficiency during literature review and evidence synthesis process, the tool does not guarantee the completeness, correctness, or reliability of its results. Users should exercise critical evaluation and cross-check all information before including it in their research.</li>
        <li><strong>Data Responsibility & Privacy:</strong> Uploaded PDFs are processed only within the session and are not stored or collected by developer. However, users remain responsible for ensuring that they have the legal and ethical right to upload and process the documents they submit.</li>
        <li><strong>Non-Liability:</strong> The developer is not liable for any direct, indirect, or consequential damages resulting from the use of this tool, including but not limited to errors in screening, extraction, data interpretation, or research outcomes.</li>
        <li><strong>Academic & Ethical Use:</strong> ReviewAid is intended solely for lawful academic and research purposes. Users must ensure compliance with all relevant institutional guidelines, copyright laws, and ethical standards.</li>
        <li><strong>AI Limitations:</strong> ReviewAid uses AI algorithms to assist with literature screening and data extraction. Outputs may contain errors, omissions, or biases, and should not be considered a substitute for expert review or professional judgment.</li>
        <li><strong>No Warranty:</strong> ReviewAid is provided "as is" without any warranty of any kind, either expressed or implied, including but not limited to accuracy, completeness, or fitness for a particular purpose.</li>
        <li><strong>Transparency & Citation:</strong> Many researchers use ReviewAid during their review process without disclosing the involvement of AI tools. Users are strongly encouraged to maintain transparency by acknowledging the use of ReviewAid in their methodology. Citation is appreciated whenever ReviewAid contributes to the research workflow, whether as a primary screener, a secondary/third-person validator, or simply as a reference tool.</li>
    </ul>
    <p>Proper attribution supports ethical research practices and helps sustain ongoing development and improvement of the tool for the academic community.</p>


        """, unsafe_allow_html=True)
    
    agree = st.checkbox("I acknowledge the disclaimer.")
    if agree and st.button("I Agree & Continue"):
        st.session_state.disclaimer_acknowledged = True
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
    st.stop()

SCREENER_API_KEY = os.getenv("SCREENER_API_KEY")
EXTRACTOR_API_KEY = os.getenv("EXTRACTOR_API_KEY")

if st.session_state.app_mode == "screener":
    if not SCREENER_API_KEY:
        st.error("Screener API key is not set. Please set the SCREENER_API_KEY environment variable.")
elif st.session_state.app_mode == "extractor":
    if not EXTRACTOR_API_KEY:
        st.error("Extractor API key is not set. Please set the EXTRACTOR_API_KEY environment variable.")

if st.session_state.app_mode is None:
    st.markdown("<div class='mode-selection'>", unsafe_allow_html=True)
    st.markdown("## Select Application Mode")
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class="mode-card">
            <div class="mode-title">Full-text Paper Screener</div>
            <div class="mode-description">
            Screen research papers based on PICO (Population, Intervention, Comparison, Outcome) criteria.
            <ul>
                <li>Define inclusion/exclusion criteria</li>
                <li>Upload PDF papers</li>
                <li>Get AI-powered screening decisions</li>
                <li>Export results</li>
            </ul>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Select Screener", key="screener_btn"):
            st.session_state.app_mode = "screener"
            st.rerun()
    
    with col2:
        st.markdown("""
        <div class="mode-card">
            <div class="mode-title">Full-text Data Extractor</div>
            <div class="mode-description">
            Extract specific data fields from research papers.
            <ul>
                <li>Define fields to extract</li>
                <li>Upload PDF papers</li>
                <li>Get AI-powered data extraction</li>
                <li>Export results</li>
            </ul>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Select Extractor", key="extractor_btn"):
            st.session_state.app_mode = "extractor"
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("""
    <div class="important-note-box">
        <strong>Note:</strong> The purpose of ReviewAid is not to substitute manual screening and data extraction but to serve as an additional, independent reference that helps minimise manual errors and improve the precision and reliability of the research process. 
          <strong>Have any Errors?</strong> Please visit the <a href="https://reviewaid.github.io/Documentation" target="_blank">Documentation section</a>.  
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<div class='citation-section'>", unsafe_allow_html=True)
    display_citation_section()
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("---")
    
    st.markdown("<div class='support-section'>", unsafe_allow_html=True)
    st.markdown("## Support")
    st.markdown("""
    <div class="support-description">
        ReviewAid is an open-source academic tool designed to streamline the literature review and evidence synthesis process. If this tool benefits your research, your support is greatly appreciated. I will be adding links to the <a href="https://github.com/aurumz-rgb/ReviewAid" target="_blank">GitHub repository</a> — please check it out to explore the source code, contribute, or support ongoing development.
    </div>
    <div class="support-description">
        I will also include my personal link to my other projects, where you can discover additional research-focused tools and resources. Check out my personal link <a href="https://aurumz-rgb.github.io" target="_blank">here</a>. If you have questions, or are interested in collaborating, feel free to reach out. I am always happy to connect with fellow researchers.
    </div>
    """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


    stats = get_firebase_stats()
    st.markdown("---")
    st.markdown("## 🌟 Global Platform Statistics")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Papers Screened", stats["papers_screened"])
    c2.metric("Papers Extracted", stats["papers_extracted"])
    c3.metric("Total Visits", stats["total_visits"])
  
    
    st.markdown("""
    <div class="mode-selection-footer">
        <div>© 2025 Vihaan Sahu – Licensed under Apache 2.0</div>
        <div class="center-text"><a href="https://github.com/aurumz-rgb/ReviewAid" target="_blank">GitHub Repository</a></div>
        <div class="right-text">Open-source</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.stop()


def extract_pdf_content(pdf_bytes):
    """
    Modified to accept raw bytes to prevent duplicate reading into memory.
    Reads file bytes ONCE, calculates hash, and passes bytes to PyMuPDF.
    """
    update_terminal_log("Initializing PDF extraction engine (PyMuPDF)...", "DEBUG")
    doc = None
    try:
       
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = doc.page_count
        update_terminal_log(f"Document opened successfully. Total pages: {page_count}", "INFO")
        
        MAX_CHAR_LIMIT = 600000 
        
        full_text_parts = []
        current_length = 0
        references_found = False
        
        for i, page in enumerate(doc):
            if current_length > MAX_CHAR_LIMIT and not references_found:
                 update_terminal_log(f"Token limit reached at Page {i+1}. Stopping PDF read.", "INFO")
                 break
            
            page_text = page.get_text()
            
            if re.search(r'(?:\n|\r\n){1,2}(References|Reference|Bibliography)(?:\s|\r?\n|$)', page_text, re.IGNORECASE):
                references_found = True 

            full_text_parts.append(page_text)
            current_length += len(page_text)
            
            del page_text 
            
        full_text = "".join(full_text_parts)
        del full_text_parts

        ref_match = re.search(r'(?:\n|\r\n){1,2}(References|Reference|Bibliography)(?:\s|\r?\n|$)', full_text, re.IGNORECASE)
        
        if ref_match:
            main_text_end = ref_match.start()
            full_text = full_text[:main_text_end]
            update_terminal_log("References section detected and removed to save tokens.", "INFO")
            del ref_match

        metadata = doc.metadata
        title = metadata.get('title', '')
        author = metadata.get('author', '')
        del metadata 
        
        update_terminal_log(f"Metadata read -> Title: '{title}', Author: '{author}'", "DEBUG")
        
        year = ''
        if not year:
            if page_count > 0:
                year_match = re.search(r'\b(19|20)\d{2}\b', full_text[:5000]) 
                if year_match:
                    year = year_match.group()
                    update_terminal_log(f"Year found on Page 1: {year}", "DEBUG")
        
        if not year:
            creation_date = doc.metadata.get('creationDate', '')
            if creation_date:
                year_match = re.search(r'\b(19|20)\d{2}\b', creation_date)
                if year_match:
                    year = year_match.group()
                    update_terminal_log(f"Year derived from creationDate: {year}", "DEBUG")
        
        return full_text, title, author, year
        
    except Exception as e:
        update_terminal_log(f"Error during PDF extraction: {str(e)}", "ERROR")
        return "", "", "", ""
    finally:
        if doc:
            doc.close()
            del doc

def preprocess_text_for_ai(text, max_tokens=MAX_INPUT_TOKENS_SCREENER):
    if "  " in text or "\n" in text:
        text = " ".join(text.split())
    
    char_limit = max_tokens * 4
    if len(text) > char_limit:
        update_terminal_log(f"Text exceeds token limit ({len(text)} > {char_limit}). Truncating...", "WARN")
        text = text[:char_limit]
    
    return text

def estimate_confidence(text, mode="screener", criteria_dict=None, extracted_data=None, fields_list=None):
    update_terminal_log(f"Calculating heuristic confidence for mode: {mode}", "DEBUG")
    
    if not text or len(text.strip()) < 30:
        return 0.1 

    text_lower = text.lower()

    if mode == "screener":
        match_count = 0
        total_criteria = 0
        

        def count_matches(criteria_string):
            nonlocal match_count, total_criteria
            if not criteria_string or not criteria_string.strip():
                return
   
            items = [c.strip() for c in criteria_string.split(",") if c.strip()]
            total_criteria += len(items)
            for item in items:
                if item.lower() in text_lower:
                    match_count += 1
        
        if criteria_dict:
            count_matches(criteria_dict.get("pop_inc", ""))
            count_matches(criteria_dict.get("pop_exc", ""))
            count_matches(criteria_dict.get("int_inc", ""))
            count_matches(criteria_dict.get("int_exc", ""))
            count_matches(criteria_dict.get("comp_inc", ""))
            count_matches(criteria_dict.get("comp_exc", ""))
            count_matches(criteria_dict.get("outcome", ""))

        if total_criteria == 0:
            update_terminal_log("No criteria provided for heuristic estimation. Defaulting to 0.4.", "DEBUG")
            return 0.4
        
        score = match_count / total_criteria
        
        if score > 0.8:
            score = min(score + 0.1, 1.0)
        
        update_terminal_log(f"Screener Heuristic: {match_count}/{total_criteria} criteria matched. Score: {score:.2f}", "DEBUG")
        return round(score, 2)

    elif mode == "extractor":
        if not extracted_data or not isinstance(extracted_data, dict):
            update_terminal_log("No extracted data available for validation. Defaulting to 0.4.", "DEBUG")
            return 0.4
            
        valid_fields = 0
        found_fields = 0
        
        for key, value in extracted_data.items():
            if value and str(value).strip() != "Not Found":
                valid_fields += 1

                val_str = str(value).strip()
                if len(val_str) > 5: 
                    if val_str.lower() in text_lower:
                        found_fields += 1
                    else:
                        words = val_str.split()[:3] 
                        if all(word in text_lower for word in words):
                            found_fields += 1
                elif len(val_str) > 0:
                     if val_str.lower() in text_lower:
                        found_fields += 1
        
        if valid_fields == 0:
            return 0.1
        
        score = found_fields / valid_fields
        update_terminal_log(f"Extractor Heuristic: {found_fields}/{valid_fields} fields verified in text. Score: {score:.2f}", "DEBUG")
        return round(score, 2)

    return 0.4

def df_from_results(results):
    rows = []
    for r in results:
        row = {
            "Filename": r.get("filename", ""),
            "Title": r.get("title", ""),
            "Author": r.get("author", ""),
            "Year": r.get("year", ""),
            "Confidence": r.get("confidence", "")
        }
        
        status = r.get("status", "").lower()
        if status == "include":
            row["Reason for Inclusion"] = r.get("reason", "")
        elif status == "exclude":
            row["Reason for Exclusion"] = r.get("reason", "")
        elif status == "maybe":
            row["Reason for Maybe"] = r.get("reason", "")
            
        rows.append(row)
    return pd.DataFrame(rows)

def df_from_extracted_results(results):
    rows = []
    for r in results:
        row = {
            "Filename": r.get("filename", ""),
            "Confidence": r.get("confidence", "")
        }
        row.update(r.get("extracted", {}))
        rows.append(row)
    return pd.DataFrame(rows)



def to_docx(df):
    """
    Generates a professional DOCX file with ReviewAid branding.
    Features:
    - Header with Logo (Right).
    - Footer with Text "Generated by ReviewAid on [Date]" (Left).
    - Styled table with bold headers and zebra striping.
    - Clean font layout (Arial).
    - "Results" heading made bolder.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    section = doc.sections[0]
    

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    logo_path = os.path.join("assets", "RA_transparent.png")
    if os.path.exists(logo_path):
        try:
          
            run_logo = header_para.add_run()
            run_logo.add_picture(logo_path, height=Inches(0.5))
        except Exception:
            pass 

  
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    

    generated_date = datetime.now().strftime("%Y-%m-%d")
    footer_text = f"Generated by ReviewAid on {generated_date}"
    
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = 'Arial'
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

  
    title = doc.add_heading('Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Shading Accent 1' 
    

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        cell = hdr_cells[i]
        cell.text = str(col)
        
        cell_par = cell.paragraphs[0]
        cell_par.runs[0].font.bold = True
        cell_par.runs[0].font.name = 'Arial'
        cell_par.runs[0].font.size = Pt(10)
        cell_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4189DC') 
        cell._element.get_or_add_tcPr().append(shading_elm)
        
  
        cell_par.runs[0].font.color.rgb = RGBColor(255, 255, 255)


    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = row_cells[i]
            para = cell.paragraphs[0]
            
         
            safe_val = str(val)
            safe_val = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', safe_val)
            
            run = para.add_run(safe_val)
            run.font.name = 'Arial'
            run.font.size = Pt(9) 
            
     
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(OxmlElement('w:vAlign'))
            tcPr[-1].set(qn('w:val'), 'top')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def to_pdf(df):
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Arial", size=8)
    

    logo_path = os.path.join("assets", "RA_transparent.png")
    if os.path.exists(logo_path):
        try:
    
            pdf.image(logo_path, x=250, y=10, w=30)
        except:
            pass 


    generated_date = datetime.now().strftime("%Y-%m-%d")
    
  
    

    col_width = 277 / len(df.columns) 
    

    pdf.set_fill_color(65, 137, 220) 
    pdf.set_text_color(255, 255, 255)
    
    for col in df.columns:
        pdf.cell(col_width, 7, str(col), border=1, align='C', fill=True)
    pdf.ln(7)
    

    pdf.set_text_color(0, 0, 0)
    pdf.set_fill_color(245, 245, 245) 
    
    is_odd = True
    for _, row in df.iterrows():
        if not is_odd:
             pdf.set_fill_color(245, 245, 245)
        else:
             pdf.set_fill_color(255, 255, 255)
             
        for val in row:
         
            clean_text = str(val).encode('latin-1', 'replace').decode('latin-1')
 
            pdf.multi_cell(col_width, 6, clean_text, border=1, fill=(not is_odd))
        pdf.ln(6)
        is_odd = not is_odd


    pdf.set_y(-15)
    pdf.set_font("Arial", size=7, style='I')
    pdf.set_text_color(128, 128, 128)
    pdf.set_x(10) 
    pdf.cell(0, 10, f"Generated by ReviewAid on {generated_date}", 0, 0, 'L')
    
    return pdf.output(dest='S').encode('latin1')

def to_csv(df):
   
    buffer = io.StringIO()
    buffer.write("# Generated by ReviewAid\n")
    df.to_csv(buffer, index=False)
    return buffer.getvalue().encode('utf-8')

def to_excel(df):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='ReviewAid Data')
        
        
        workbook = writer.book
        worksheet = writer.sheets['ReviewAid Data']
        

        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#4189DC', 
            'font_color': '#FFFFFF',
            'border': 1
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            

        generated_date = datetime.now().strftime("%Y-%m-%d")

        worksheet.set_footer(f'&LGenerated by ReviewAid on {generated_date}')
        
    buffer.seek(0)
    return buffer.getvalue()


def find_exclusion_matches(text, exclusion_lists):
    matches = []
    update_terminal_log("Scanning text for exclusion keywords...", "DEBUG")
    for criteria in exclusion_lists:
        criteria = criteria.strip()
        if criteria:
            if criteria.lower() in text.lower():
                update_terminal_log(f"Match found for exclusion criteria: '{criteria}'", "INFO")
                matches.append(criteria)
            else:
                update_terminal_log(f"No match for exclusion criteria: '{criteria}'", "DEBUG")
    return matches

def update_terminal_log(msg, level="INFO"):
    timestamp = datetime.now().strftime("%H:%M:%S.%f")[:-3] 
    
    if level == "INFO":
        color_class = "log-info"
    elif level == "SUCCESS":
        color_class = "log-success"
    elif level == "WARN":
        color_class = "log-warn"
    elif level == "ERROR":
        color_class = "log-error"
    elif level == "SYSTEM":
        color_class = "log-system"
    elif level == "DEBUG":
        color_class = "log-debug"
    else:
        color_class = "log-info"
        
    log_entry = f'<span class="terminal-line"><span class="terminal-timestamp">[{timestamp}]</span><span class="{color_class}">{level}</span>: {html.escape(msg)}</span>'
    
    st.session_state.terminal_logs.append(log_entry)
    
    if len(st.session_state.terminal_logs) > MAX_LOG_ENTRIES:
        st.session_state.terminal_logs.pop(0)
    
    full_log_html = '<div class="terminal-container" id="terminal-container">' + "".join(st.session_state.terminal_logs) + '</div>'
    

    scroll_script = """
    <script>
        var element = document.getElementById("terminal-container");
        if(element){
            // Force a slight delay to ensure render is complete
            setTimeout(function() {
                element.scrollTop = element.scrollHeight;
            }, 50);
        }
    </script>
    """
    
    st.session_state.terminal_placeholder.markdown(full_log_html + scroll_script, unsafe_allow_html=True)


if st.session_state.app_mode == "screener":
    st.markdown("## Full-text Paper Screener")
    
    st.subheader("Population Criteria")
    population_inclusion = st.text_area("Population Inclusion Criteria", placeholder="e.g. Adults aged 18–65 with MS")
    population_exclusion = st.text_area("Population Exclusion Criteria", placeholder="e.g. Patients with comorbid autoimmune diseases")

    st.subheader("Intervention Criteria")
    intervention_inclusion = st.text_area("Intervention Inclusion Criteria", placeholder="e.g. Natalizumab treatment ≥ 6 months")
    intervention_exclusion = st.text_area("Intervention Exclusion Criteria", placeholder="e.g. Dose outside approved range")

    st.subheader("Comparison Criteria")
    comparison_inclusion = st.text_area("Comparison Inclusion Criteria", placeholder="e.g. Placebo or no treatment")
    comparison_exclusion = st.text_area("Comparison Exclusion Criteria", placeholder="e.g. Active comparator like interferon beta")

    st.subheader("Outcome Criteria (Optional)")
    outcome_criteria = st.text_area("Outcome Criteria", placeholder="e.g. Annualized relapse rate, disability progression")

    uploaded_pdfs = st.file_uploader("Upload PDF Files", accept_multiple_files=True)
    
    
    fields_list = []

elif st.session_state.app_mode == "extractor":
    st.markdown("## Full-text Data Extractor")
    fields = st.text_input("Fields to Extract (comma-separated)", placeholder="e.g. Author, Year, Study Design, Sample Size, Conclusion")
    fields_list = [f.strip() for f in fields.split(",") if f.strip()]
    
    if "Paper Title" not in fields_list:
        fields_list.insert(0, "Paper Title")
    
    if len(fields_list) == 1 and fields_list[0] == "Paper Title":
        st.info("If left Empty, Only Paper Title will be extracted. Add more fields to extract additional information.")
    
    

    uploaded_pdfs = st.file_uploader("Upload PDF Files", accept_multiple_files=True)
    
    population_inclusion = ""
    population_exclusion = ""
    intervention_inclusion = ""
    intervention_exclusion = ""
    comparison_inclusion = ""
    comparison_exclusion = ""
    outcome_criteria = ""

if 'terminal_placeholder' not in st.session_state:
    st.session_state.terminal_placeholder = st.empty()

if st.button("Process Papers" if st.session_state.app_mode == "extractor" else "Screen Papers"):
   
    st.session_state.included_results = []
    st.session_state.excluded_results = []
    st.session_state.maybe_results = []
    st.session_state.extracted_results = []
    

    st.session_state.batch_file_hashes = {}

    if not uploaded_pdfs:
        st.warning("Please upload at least one PDF file.")
        st.stop()
    
    if st.session_state.app_mode == "screener" and not any([
        population_inclusion.strip(), population_exclusion.strip(),
        intervention_inclusion.strip(), intervention_exclusion.strip(),
        comparison_inclusion.strip(), comparison_exclusion.strip(),
        outcome_criteria.strip()
    ]):
        st.warning("Please enter at least one inclusion or exclusion criterion.")
        st.stop()

    if st.session_state.app_mode == "screener":
        if not SCREENER_API_KEY:
            st.error("Screener API key is not set. Please set the SCREENER_API_KEY environment variable.")
            st.stop()
        api_key = SCREENER_API_KEY
    else:
        if not EXTRACTOR_API_KEY:
            st.error("Extractor API key is not set. Please set the EXTRACTOR_API_KEY environment variable.")
            st.stop()
        api_key = EXTRACTOR_API_KEY

    st.session_state.terminal_logs = []
    with st.expander("System Terminal (Background Processing)", expanded=True):
        st.session_state.terminal_placeholder = st.empty()
        update_terminal_log("Initializing processing session...", "SYSTEM")
        update_terminal_log(f"Mode detected: {st.session_state.app_mode}", "INFO")
        update_terminal_log(f"Files to process: {min(len(uploaded_pdfs), 2000)}", "INFO")
        update_terminal_log("Allocating resources...", "DEBUG")

    max_papers = 2000
    total_pdfs = min(len(uploaded_pdfs), max_papers)
    
    
    status_placeholder = st.empty()
    progress_bar = st.progress(0)
    
    papers_processed_in_batch = 0

    
    try:
        for idx, pdf in enumerate(uploaded_pdfs[:max_papers], 1):
 
            gc.collect()
            
            update_terminal_log(f"--- Starting File {idx}/{total_pdfs}: {pdf.name} ---", "SYSTEM")
            
            try:
                start_time_file = time.time()

     
                pdf.seek(0)
                pdf_bytes = pdf.read()
                pdf_hash = hashlib.sha256(pdf_bytes).hexdigest()
                
                if pdf_hash in st.session_state.batch_file_hashes:
                    update_terminal_log(f"Duplicate file detected (Hash match). Using cached result.", "WARN")
                    cached_result = st.session_state.batch_file_hashes[pdf_hash]
        
                    cached_result["filename"] = pdf.name
                    
                    if st.session_state.app_mode == "screener":
                        status = cached_result.get("status", "").lower()
                        if "include" in status:
                            st.session_state.included_results.append(cached_result)
                        elif "exclude" in status:
                            st.session_state.excluded_results.append(cached_result)
                        elif "maybe" in status:
                            st.session_state.maybe_results.append(cached_result)
                        else:
                            st.session_state.excluded_results.append(cached_result)
                        update_processing_stats("screener", 1)
                    else:
                        st.session_state.extracted_results.append(cached_result)
                        update_processing_stats("extractor", 1)
                    
                    papers_processed_in_batch += 1
                    
                 
                    percent = int((idx / total_pdfs) * 100)
                    status_placeholder.markdown(f"<h4 style='text-align: center; color: #4189DC;'>{percent}% Work Done... Processing <span style='color: white'>{pdf.name}</span> (Cached)</h4>", unsafe_allow_html=True)
                    progress_bar.progress(idx / total_pdfs)
                    update_terminal_log("Skipped API call. Using cached data.", "SUCCESS")
                    
         
                    del pdf_bytes
                    continue
                else:
                    update_terminal_log("New file detected. Proceeding with extraction.", "DEBUG")
                

                text, title, author, year = extract_pdf_content(pdf_bytes)
                
         
                del pdf_bytes 

                if not text.strip():
                    update_terminal_log(f"PDF '{pdf.name}' appears empty or unreadable.", "WARN")
                    update_terminal_log("Skipping this file.", "WARN")
                    progress_bar.progress(idx / total_pdfs)
                    continue
                
                update_terminal_log(f"Text extracted. Length: {len(text)} chars.", "INFO")

                full_text_backup = text

                if st.session_state.app_mode == "screener":
                    text = preprocess_text_for_ai(text, max_tokens=MAX_INPUT_TOKENS_SCREENER)
                else:
                    text = preprocess_text_for_ai(text, max_tokens=MAX_INPUT_TOKENS_EXTRACTOR)

                if st.session_state.app_mode == "screener":
                    update_terminal_log(f"Text preprocessed. Input Tokens: ~{MAX_INPUT_TOKENS_SCREENER}.", "INFO")
                else:
                    update_terminal_log(f"Text preprocessed. Input Tokens: ~{MAX_INPUT_TOKENS_EXTRACTOR}.", "INFO")
 
                criteria_dict = {
                    "pop_inc": population_inclusion,
                    "pop_exc": population_exclusion,
                    "int_inc": intervention_inclusion,
                    "int_exc": intervention_exclusion,
                    "comp_inc": comparison_inclusion,
                    "comp_exc": comparison_exclusion,
                    "outcome": outcome_criteria
                }
            
                confidence = estimate_confidence(
                    text, 
                    mode=st.session_state.app_mode, 
                    criteria_dict=criteria_dict, 
                    extracted_data=None, 
                    fields_list=fields_list
                )
                update_terminal_log(f"Initial heuristic confidence score estimated: {confidence}", "INFO")

                if st.session_state.app_mode == "screener":
                    all_exclusions = []
                    for block in [population_exclusion, intervention_exclusion, comparison_exclusion]:
                        if block.strip():
                            all_exclusions.extend([c.strip() for c in block.split(",") if c.strip()])
                    
                    update_terminal_log(f"Total exclusion criteria loaded: {len(all_exclusions)}", "INFO")
                    matches_exc = find_exclusion_matches(text, all_exclusions)

                    all_inclusions = []
                    for block in [population_inclusion, intervention_inclusion, comparison_inclusion]:
                        if block.strip():
                            all_inclusions.extend([c.strip() for c in block.split(",") if c.strip()])
                    
                    matches_inc = []
                    update_terminal_log("Scanning text for inclusion keywords...", "DEBUG")
                    for criteria in all_inclusions:
                        if criteria.strip() and criteria.lower() in text.lower():
                            update_terminal_log(f"Match found for inclusion criteria: '{criteria}'", "INFO")
                            matches_inc.append(criteria)
                
            
                    del all_exclusions, all_inclusions, matches_inc

                    if len(matches_exc) >= 1 and len(matches_inc) == 0:
                        exclusion_reason = (
                            f"Auto-excluded because {len(matches_exc)} exclusion criteria matched: {', '.join(matches_exc)}"
                        )
                        confidence = 1.0 
                        
                        result = {
                            "filename": pdf.name,
                            "status": "Exclude",
                            "reason": exclusion_reason[:500], 
                            "confidence": confidence,
                            "title": title,
                            "author": author,
                            "year": year
                        }
                        st.session_state.batch_file_hashes[pdf_hash] = result
                        st.session_state.excluded_results.append(result)
                        update_terminal_log(f"Found {len(matches_exc)} exclusion matches: {matches_exc}", "WARN")
                        update_terminal_log(f"Result: EXCLUDED (Auto-rule) with Confidence 1.0", "WARN")
                        
      
                        percent = int((idx / total_pdfs) * 100)
                        status_placeholder.markdown(f"<h4 style='text-align: center; color: #4189DC;'>{percent}% Work Done... Processing <span style='color: white'>{pdf.name}</span></h4>", unsafe_allow_html=True)
                        progress_bar.progress(idx / total_pdfs)
                        
                        update_processing_stats("screener", 1)
                        papers_processed_in_batch += 1

              
                        del text, full_text_backup, matches_exc, exclusion_reason
                        continue
                    
                    elif len(matches_exc) >= 1 and len(matches_inc) >= 1:
                        update_terminal_log("Conflict detected: Both Exclusion and Inclusion keywords found. Sending to AI for resolution.", "WARN")
                 
                    update_terminal_log("Constructing PICO prompt for AI...", "INFO")
                    prompt = f"""
You are an expert systematic reviewer. Your task is to screen a research paper based on specific PICO criteria.

**CRITICAL INSTRUCTION:**
Return your response as a SINGLE valid JSON object. Do not include markdown formatting (like ```json), do not add comments, and do not include conversational filler text.

**Population**
Inclusion: {population_inclusion}
Exclusion: {population_exclusion}

**Intervention**
Inclusion: {intervention_inclusion}
Exclusion: {intervention_exclusion}

**Comparison**
Inclusion: {comparison_inclusion}
Exclusion: {comparison_exclusion}

**Outcomes**: {outcome_criteria}

**Paper Text:**
\"\"\"
{text}
\"\"\"

**Task:**
1. Classify the paper as "Include", "Exclude", or "Maybe" based strictly on the criteria.
2. Provide a detailed reason for the classification.
3. Extract the Paper Title, Main Author, and Publication Year.
4. If a value is not found, use "Not Found".
5. **CONFIDENCE SCORE**: Rate your confidence (0.0 to 1.0). 
   - 1.0 = The paper perfectly matches or perfectly violates the criteria with explicit evidence.
   - 0.8 - 0.9 = High confidence based on strong evidence.
   - 0.5 - 0.7 = Moderate confidence (Some ambiguity in criteria or text).
   - < 0.5 = Low confidence (Guessing, criteria vague, or text unclear).

**JSON Format Required:**
{{
  "status": "Include",
  "reason": "Detailed classification reason explaining why it fits or fails the criteria.",
  "title": "Full paper title extracted from text",
  "author": "Main author name",
  "year": "2023",
  "confidence": 0.95
}}
"""
                else:
                    update_terminal_log(f"Preparing extraction fields: {', '.join(fields_list)}", "INFO")
                    field_descriptions = {
                        "Paper Title": "The full title of the research paper",
                        "Author": "The main author(s) of the paper",
                        "Year": "The publication year of the paper",
                        "Journal": "The journal where the paper was published",
                        "DOI": "The Digital Object Identifier of the paper",
                        "Abstract": "A brief summary of the paper's content",
                        "Keywords": "Key terms associated with the paper",
                        "Study Design": "The methodology used in the study (e.g., randomized controlled trial, cohort study)",
                        "Sample Size": "The number of participants in the study",
                        "Intervention": "The treatment or intervention being studied",
                        "Comparison": "The control or comparison group",
                        "Outcome": "The main results or findings of the study",
                        "Conclusion": "The authors' conclusion based on the findings",
                        "Funding": "Information about who funded the research",
                        "Conflicts of Interest": "Any declared conflicts of interest by the authors"
                    }
                    
                    prompt = "Extract the following information from the research paper:\n\n"
                    for field in fields_list:
                        description = field_descriptions.get(field, f"Information about {field}")
                        prompt += f"- {field}: {description}\n"
                    
                    prompt += f"""
**Paper Text:**
\"\"\"
{text}
\"\"\"

**CRITICAL INSTRUCTION:**
Return your response as a SINGLE valid JSON object. Do not include markdown formatting. Ensure all keys are present.
If a field is not found in the text, use the value "Not Found".
**CONFIDENCE SCORE**: Rate your confidence (0.0 to 1.0).
- 1.0 = All extracted fields are explicitly stated in the text.
- 0.8 - 0.9 = Most fields are explicit, some inferred.
- 0.5 - 0.7 = Some fields missing or ambiguous.
- < 0.5 = Data largely missing or garbled.

**JSON Format Required:**
{{
  "extracted": {{
"""
                    for field in fields_list:
                        prompt += f'    "{field}": "",\n'
                    
                    prompt = prompt.rstrip(",\n") + "\n  },\n"
                    prompt += '  "confidence": 0.0\n}'
                    prompt += "\nEnsure that JSON is valid. Use 'Not Found' for missing data.\n"
                
            
                raw_result = None
            
                processing_successful = False

               
                max_api_attempts = 3 
                retries_per_api_attempt = 3
                
                for api_attempt in range(max_api_attempts):
                    if api_attempt > 0:
                        update_terminal_log("Previous attempt failed after retries. Initiating NEW API call for this paper...", "WARN")
                    
                    for retry_idx in range(retries_per_api_attempt):
                        if st.session_state.app_mode == "extractor":
                            raw_result = query_zai(prompt, api_key, temperature=0.1, max_tokens=MAX_OUTPUT_TOKENS)
                        else:
                            raw_result = query_zai(prompt, api_key, temperature=0.1, max_tokens=2048)

       
                        if raw_result is None:
                            update_terminal_log("API returned None after exhaustive retries. Falling back to Regex.", "ERROR")
                            processing_successful = False
                            break 
                        
                        if raw_result and raw_result.strip():
                            processing_successful = True
                            break 
                        
                        if retry_idx < retries_per_api_attempt - 1:
                             update_terminal_log(f"API returned empty response. Retry {retry_idx + 2}/{retries_per_api_attempt}...", "ERROR")
                             time.sleep(2)
                    
                    if processing_successful:
                        break
                
              
                del prompt 

                if not processing_successful:
                    update_terminal_log("AI processing failed. Switching to Regex Fallback (Local Processing) to ensure result is generated.", "WARN")

                if raw_result:
                     update_terminal_log(f"API response received. Length: {len(raw_result)} chars.", "SUCCESS")
                     update_terminal_log("Parsing JSON response...", "INFO")
                else:
                     update_terminal_log("Using Fallback Strategy for parsing...", "WARN")

                if st.session_state.app_mode == "screener":
                    result = parse_result(raw_result, api_key, mode="screener", original_text=full_text_backup, fields_list=fields_list)
                else:
                    result = parse_result(raw_result, api_key, mode="extractor", fields_list=fields_list, original_text=full_text_backup)
                    
                    if "extracted" not in result:
                        result["extracted"] = {}
                    
                    for field in fields_list:
                        if field not in result["extracted"]:
                            result["extracted"][field] = "Not Found"
                
     
                if result:
                    processing_successful = True
                    
                    if "filename" not in result:
                        result["filename"] = pdf.name
                
           
                if raw_result:
                    del raw_result
                
      
                del full_text_backup
                gc.collect()

                if result and processing_successful:
                    update_terminal_log("Result generation completed (AI or Fallback).", "SUCCESS")

                    ai_confidence = result.get("confidence", None)
                    if ai_confidence is not None:
                        try:
                            confidence = float(ai_confidence)
                            confidence = max(0.0, min(1.0, confidence))
                            update_terminal_log(f"Using reported confidence: {confidence}", "INFO")
                        except ValueError:
                            update_terminal_log(f"Confidence invalid format. Using fallback.", "WARN")

                            confidence = estimate_confidence(
                                text, 
                                mode=st.session_state.app_mode, 
                                criteria_dict=criteria_dict, 
                                extracted_data=result.get("extracted") if st.session_state.app_mode == "extractor" else None,
                                fields_list=fields_list
                            )
                    else:
                        update_terminal_log(f"Confidence missing. Using heuristic fallback.", "WARN")

                        confidence = estimate_confidence(
                            text, 
                            mode=st.session_state.app_mode, 
                            criteria_dict=criteria_dict, 
                            extracted_data=result.get("extracted") if st.session_state.app_mode == "extractor" else None,
                            fields_list=fields_list
                        )

                    result["confidence"] = confidence
                    
                    if confidence < 0.5:
                        result["flags"] = ["low_confidence"]
                        update_terminal_log("Flagged: Low confidence score.", "WARN")

                    if st.session_state.app_mode == "screener":
                        if "title" not in result or not result["title"] or result["title"] == "":
                            result["title"] = title
                        if "author" not in result or not result["author"] or result["author"] == "":
                            result["author"] = author
                        if "year" not in result or not result["year"] or result["year"] == "":
                            result["year"] = year
                    
                    del title, author, year, text 

                    if st.session_state.app_mode == "screener":
                        status = result.get("status", "").strip().lower() 
                        if "include" in status:
                            status = "include"
                        elif "exclude" in status:
                            status = "exclude"
                        elif "maybe" in status:
                            status = "maybe"
                        else:
                             status = "exclude"
                        
                        if len(result.get("reason", "")) > 500:
                            result["reason"] = result["reason"][:500] + "..."

                        if status == "include":
                            st.session_state.included_results.append(result)
                            update_terminal_log(f"Final Decision: INCLUDE", "SUCCESS")
                        elif status == "exclude":
                            st.session_state.excluded_results.append(result)
                            update_terminal_log(f"Final Decision: EXCLUDE", "SUCCESS")
                        elif status == "maybe":
                            st.session_state.maybe_results.append(result)
                            update_terminal_log(f"Final Decision: MAYBE", "SUCCESS")
                        else:
                            st.session_state.excluded_results.append(result)
                            update_terminal_log(f"Final Decision: EXCLUDE (Default)", "INFO")
                        
                        update_processing_stats("screener", 1)
                        papers_processed_in_batch += 1

                    else:
                        for key in result["extracted"]:
                            if isinstance(result["extracted"][key], str) and len(result["extracted"][key]) > 1000:
                                result["extracted"][key] = result["extracted"][key][:1000] + "..."

                        st.session_state.extracted_results.append(result)
                        update_terminal_log("Data extraction completed.", "SUCCESS")
                        update_processing_stats("extractor", 1)
                        papers_processed_in_batch += 1
            

                    st.session_state.batch_file_hashes[pdf_hash] = result
                
                elapsed = time.time() - start_time_file
                update_terminal_log(f"File processed in {elapsed:.2f}s.", "SYSTEM")

      
            except Exception as e:
                update_terminal_log(f"CRITICAL ERROR processing {pdf.name}: {str(e)}", "ERROR")
                import traceback
                update_terminal_log(f"Traceback: {traceback.format_exc()}", "ERROR")
              
                gc.collect()
            
            
            percent = int((idx / total_pdfs) * 100)
            status_placeholder.markdown(f"<h6 style='text-align: center; color: 'white';'>Processed {percent}% of papers... Processing <span style='color: #4189DC'> {pdf.name}</span></h6>", unsafe_allow_html=True)
            progress_bar.progress(idx / total_pdfs)
            
        
    except Exception as e:
        update_terminal_log(f"CRITICAL BATCH ERROR: {str(e)}", "ERROR")
        import traceback
        update_terminal_log(f"Traceback: {traceback.format_exc()}", "ERROR")


    status_placeholder.empty()
    update_terminal_log("=== Batch processing complete ===", "SYSTEM")
    update_terminal_log(f"Processed {papers_processed_in_batch} papers in this session.", "SUCCESS")

if st.session_state.app_mode == "screener":
    included = len(st.session_state.included_results)
    excluded = len(st.session_state.excluded_results)
    maybe = len(st.session_state.maybe_results)
    total = included + excluded + maybe

    session_duration = time.time() - st.session_state.start_time
    avg_speed = session_duration / total if total > 0 else 0

    with st.expander("Screening Dashboard", expanded=False):
         st.metric("Papers Screened", total)
         

 
       
         colors_map = {
             "Included": "#3fb950", 
             "Excluded": "#f85149",
             "Maybe": "#d29922"     
         }
         
         fig = px.pie(
            names=["Included", "Excluded", "Maybe"],
            values=[included, excluded, maybe],
            title="Screening Decisions",
            hole=0.4,
            color=["Included", "Excluded", "Maybe"],
            color_discrete_map=colors_map
        )
         
         fig.update_layout(
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            font=dict(color='#F0F4F8'),
            legend=dict(orientation="h", yanchor="bottom", y=-0.1, xanchor="center", x=0.5),
            title_font=dict(size=20)
         )
         
         col_g1, col_g2 = st.columns([3, 1])
         with col_g1:
             st.plotly_chart(fig, width='stretch')
         with col_g2:
             html_string = fig.to_html()
             st.download_button(
                label="Download Graph",
                data=html_string,
                file_name="screening_decisions_graph.html",
                mime="text/html"
            )
             del html_string

    included_results = st.session_state.included_results
    excluded_results = st.session_state.excluded_results
    maybe_results = st.session_state.maybe_results 

    def color_status(val):
        if val == "Include":
            return 'background-color: rgba(16, 185, 129, 0.15); color: #34d399; font-weight: bold;'
        elif val == "Exclude":
            return 'background-color: rgba(239, 68, 68, 0.15); color: #f87171; font-weight: bold;'
        elif val == "Maybe":
            return 'background-color: rgba(245, 158, 11, 0.15); color: #fbbf24; font-weight: bold;'
        return ''

    if included_results:
        st.header("Included Papers")
        df_inc = df_from_results(included_results)
        styled_inc = df_inc.style.set_properties(**{'text-align': 'left'})
        st.dataframe(styled_inc, width='stretch', height=400)
        del df_inc
    else:
        st.info("No Included papers found.")

    if excluded_results:
        st.header("Excluded Papers")
        df_exc = df_from_results(excluded_results)
        styled_exc = df_exc.style.set_properties(**{'text-align': 'left'})
        st.dataframe(styled_exc, width='stretch', height=400)
        del df_exc
    else:
        st.info("No Excluded papers found.")

    if maybe_results:
        st.header("Maybe Papers")
        df_maybe = df_from_results(maybe_results)
        styled_maybe = df_maybe.style.set_properties(**{'text-align': 'left'})
        st.dataframe(styled_maybe, width='stretch', height=400)
        del df_maybe
    else:
        st.info("No Maybe papers found.")

    if included_results or excluded_results or maybe_results:
        st.markdown("""
        <div class="confidence-table-container">
            <h3 style="margin-top:0;">Confidence Score Interpretation</h3>
            <table style="width:100%; border-collapse: collapse; color: #F0F4F8;">
                <thead>
                    <tr>
                        <th>Confidence Score</th>
                        <th>Classification</th>
                        <th>Description</th>
                        <th>Implication</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>1.0 (100%)</strong></td>
                        <td>Definitive Match</td>
                        <td>Deterministic rule-based classification / No ambiguity.</td>
                        <td>Fully automated decision</td>
                    </tr>
                    <tr>
                        <td><strong>0.8 – 1.0</strong></td>
                        <td>Very High Confidence</td>
                        <td>AI strongly validates decision using explicit textual evidence.</td>
                        <td>Safe to accept</td>
                    </tr>
                    <tr>
                        <td><strong>0.6 – 0.79</strong></td>
                        <td>High Confidence</td>
                        <td>Criteria appear satisfied based on standard academic structure and content.</td>
                        <td>Review optional</td>
                    </tr>
                    <tr>
                        <td><strong>0.4 – 0.59</strong></td>
                        <td>Moderate Confidence</td>
                        <td>Ambiguous context or loosely met criteria.</td>
                        <td>Manual verification recommended</td>
                    </tr>
                    <tr>
                        <td><strong>0.1 – 0.39</strong></td>
                        <td>Low Confidence</td>
                        <td>Based mainly on heuristic keyword estimation.</td>
                        <td>High risk of error</td>
                    </tr>
                    <tr>
                        <td><strong>&lt; 0.1</strong></td>
                        <td>Unreliable</td>
                        <td>Derived from fallback or failed extraction methods.</td>
                        <td>Mandatory manual review</td>
                    </tr>
                </tbody>
            </table>
        </div>
        """, unsafe_allow_html=True)

        st.header("Export Results")

        def export_buttons(df, label_prefix):
            formats = {
                "DOCX": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", to_docx(df), f"{label_prefix.lower()}_papers.docx"),
                "CSV":  ("text/csv", to_csv(df), f"{label_prefix.lower()}_papers.csv"),
                "XLSX": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", to_excel(df), f"{label_prefix.lower()}_papers.xlsx")
            }

            for fmt, (mime, data, filename) in formats.items():
                st.download_button(
                    label=f"Download {label_prefix} as {fmt}",
                    data=data,
                    file_name=filename,
                    mime=mime
                )

        if included_results:
            st.subheader("Included Papers")
            df_inc = df_from_results(included_results)
            export_buttons(df_inc, "Included")
            del df_inc

        if excluded_results:
            st.subheader("Excluded Papers")
            df_exc = df_from_results(excluded_results)
            export_buttons(df_exc, "Excluded")
            del df_exc

        if maybe_results:
            st.subheader("Maybe Papers")
            df_maybe = df_from_results(maybe_results)
            export_buttons(df_maybe, "Maybe")
            del df_maybe

else:
    extracted_results = st.session_state.extracted_results
    total = len(extracted_results)
    
    with st.expander("Extraction Dashboard", expanded=False):
         st.metric("Papers Processed", total)
         

        
   

    if extracted_results:
        st.header("Extracted Data")
        df_ext = df_from_extracted_results(extracted_results)
        
        def style_extractor_df(df):
            return df.style.set_properties(**{'text-align': 'left'})
            
        styled_ext = style_extractor_df(df_ext)
        st.dataframe(styled_ext, width='stretch', height=400)

        st.markdown("""
        <div class="confidence-table-container">
            <h3 style="margin-top:0;">Confidence Score Interpretation</h3>
            <table style="width:100%; border-collapse: collapse; color: #F0F4F8;">
                <thead>
                    <tr>
                        <th>Confidence Score</th>
                        <th>Classification</th>
                        <th>Description</th>
                        <th>Implication</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>0.8 – 1.0</strong></td>
                        <td>Very High Confidence</td>
                        <td>AI strongly validates decision using explicit textual evidence.</td>
                        <td>Safe to accept</td>
                    </tr>
                    <tr>
                        <td><strong>0.6 – 0.79</strong></td>
                        <td>High Confidence</td>
                        <td>Criteria appear satisfied based on standard academic structure and content.</td>
                        <td>Review optional</td>
                    </tr>
                    <tr>
                        <td><strong>0.4 – 0.59</strong></td>
                        <td>Moderate Confidence</td>
                        <td>Ambiguous context or loosely met criteria.</td>
                        <td>Manual verification recommended</td>
                    </tr>
                    <tr>
                        <td><strong>0.1 – 0.39</strong></td>
                        <td>Low Confidence</td>
                        <td>Based mainly on heuristic keyword estimation.</td>
                        <td>High risk of error</td>
                    </tr>
                    <tr>
                        <td><strong>&lt; 0.1</strong></td>
                        <td>Unreliable</td>
                        <td>Derived from fallback or failed extraction methods.</td>
                        <td>Mandatory manual review</td>
                    </tr>
                </tbody>
            </table>
        </div>
        """, unsafe_allow_html=True)
        
        st.header("Export Results")
        
        def export_buttons(df, label_prefix):
            formats = {
                "DOCX": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", to_docx(df), f"{label_prefix.lower()}_data.docx"),
                "CSV":  ("text/csv", to_csv(df), f"{label_prefix.lower()}_data.csv"),
                "XLSX": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", to_excel(df), f"{label_prefix.lower()}_data.xlsx")
            }

            for fmt, (mime, data, filename) in formats.items():
                st.download_button(
                    label=f"Download {label_prefix} as {fmt}",
                    data=data,
                    file_name=filename,
                    mime=mime
                )
        
        export_buttons(df_ext, "Extracted")
        del df_ext
    else:
        st.info("No extracted data available. Upload and process papers to see results here.")

st.markdown(
    """
    <style>
        .github-link {
            display: inline-block;
            color: #4189DC;
            
            font-size: 16px;
            text-decoration: underline;
            cursor: pointer;
            transition: color 0.3s ease;
        }
        .github-link:hover {
            color: #174D8B;
        }
    </style>
   
    """,
    unsafe_allow_html=True
)

st.markdown("</div>", unsafe_allow_html=True)

st.markdown("""
    <div class="important-note-box">
        <strong>Note:</strong> The purpose of ReviewAid is not to substitute manual screening and data extraction but to serve as an additional, independent reference that helps minimise manual errors and improve the precision and reliability of the research process. 
          <strong>Have any Errors?</strong> Please visit the <a href="https://reviewaid.github.io/Documentation" target="_blank">Documentation section</a>.  
    </div>
    """, unsafe_allow_html=True)

if st.session_state.app_mode is not None:
    st.markdown("<div class='citation-section'>", unsafe_allow_html=True)
    display_citation_section()
    st.markdown("</div>", unsafe_allow_html=True)

version = "2.0.0"
try:
    current_file = __file__
    last_modified_timestamp = os.path.getmtime(current_file)
    last_updated = datetime.fromtimestamp(last_modified_timestamp).strftime("%Y-%m-%d %H:%M:%S")
except Exception:
    last_updated = "Unknown"

st.markdown(
    f"""
    <style>
    .custom-footer-container {{
        width: 100%;
        font-family: 'Times New Roman', Times, serif;
        font-size: 13.5px;
        color: #F0F4F8;
        opacity: 0.9;
        padding: 8px 20px;
        position: fixed;
        bottom: 0;
        left: 0;
        background-color: rgba(31, 41, 55, 1);
        display: flex;
        justify-content: space-between;
        align-items: center;
        z-index: 9999;
        backdrop-filter: blur(5px);
    }}
    
  
    @media screen and (max-width: 600px) {{
        .custom-footer-container {{
            flex-direction: column;
            height: auto;
            padding: 10px;
            gap: 5px;
            text-align: center;
            font-size: 11px;
        }}
    }}
    </style>

    <div class="custom-footer-container">
        <div style="white-space: nowrap; letter-spacing: 1px;">
            Made with 💙 by its Creator.
        </div>
        <div style="white-space: nowrap; letter-spacing: 1px;">
            Version {version}  
        </div>
    </div>
    """,
    unsafe_allow_html=True
)
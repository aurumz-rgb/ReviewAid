import streamlit as st
import fitz  
import time
import os 
import re
import json
import json5
import pandas as pd
import io
import html
import hashlib
import gc
from abc import ABC, abstractmethod
from typing import List, Dict, Any
from datetime import datetime
from dotenv import load_dotenv
import firebase_admin
from firebase_admin import credentials, firestore
from cryptography.fernet import Fernet
from fpdf import FPDF
from streamlit_lottie import st_lottie
import streamlit.components.v1 as components
import base64


try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from anthropic import Anthropic
except ImportError:
    Anthropic = None

try:
    import cohere
except ImportError:
    cohere = None

try:
    import ollama
except ImportError:
    ollama = None

try:
    from zai import ZaiClient
except ImportError:
    ZaiClient = None

load_dotenv()


db = None 

try:
    if not firebase_admin._apps:
        firebase_key_str = os.getenv("FIREBASE_KEY")
        if firebase_key_str:
            cred_dict = json.loads(firebase_key_str)
            cred = credentials.Certificate(cred_dict)
            firebase_admin.initialize_app(cred)
    
    if firebase_admin._apps:
        db = firestore.client()
except Exception as e:
   
    try:
        st.warning(f"Error initializing Firebase: {e}")
    except:
        print(f"Error initializing Firebase: {e}")

MAX_LOG_ENTRIES = 200
MAX_INPUT_TOKENS_SCREENER = 128000 
MAX_INPUT_TOKENS_EXTRACTOR = 128000 
MAX_OUTPUT_TOKENS = 8192


def get_firebase_stats():
    if db is None:
        return {"papers_screened": 0, "papers_extracted": 0, "total_visits": 0}
    try:
        doc_ref = db.collection("ReviewAidAnalytics").document("counters")
        doc = doc_ref.get()
        
        if doc:
            return {
                "papers_screened": doc.to_dict().get("papers_screened", 0),
                "papers_extracted": doc.to_dict().get("papers_extracted", 0),
                "total_visits": doc.to_dict().get("total_visits", 0)
            }
        else:
            
            return {"papers_screened": 0, "papers_extracted": 0, "total_visits": 0}
    except Exception:
     
        return {"papers_screened": 0, "papers_extracted": 0, "total_visits": 0}

def display_citation_section():
    """
    Displays the citation section.
    Checks session state to prevent duplicate element IDs on re-runs.
    """


    st.markdown("---")
    st.markdown("## Citation")
    st.markdown("You can check ReviewAid's preprint paper here: [ReviewAid MetaArXiV](https://osf.io/preprints/metaarxiv/3vhmt). If you are to use ReviewAid, please don't forget to cite the tool:")

    apa_citation = (
        "Sahu, V. (2025). ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0). "
        "Zenodo. https://doi.org/10.5281/zenodo.18060972"
    )

    harvard_citation = (
        "Sahu, V., 2025. ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0). "
        "Zenodo. Available at: https://doi.org/10.5281/zenodo.18060972"
    )

    mla_citation = (
        "Sahu, Vihaan. \"ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0).\" "
        "2025, Zenodo, https://doi.org/10.5281/zenodo.18060972."
    )

    chicago_citation = (
        "Sahu, Vihaan. 2025. \"ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0).\" "
        "Zenodo. https://doi.org/10.5281/zenodo.18060972."
    )

    ieee_citation = (
        "V. Sahu, \"ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0),\" "
        "Zenodo, 2025. doi: 10.5281/zenodo.18060972."
    )

    vancouver_citation = (
        "Sahu V. ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0). "
        "Zenodo. 2025. doi:10.5281/zenodo.18060972"
    )

    ris_data = """TY  - JOUR
AU  - Sahu, V
TI  - ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0)
PY  - 2025
DO  - 10.5281/zenodo.18060972
ER  -"""

    bib_data = """@misc{Sahu2025,
  author={Sahu, V.},
  title={ReviewAid: AI-Driven Full-Text Screening and Data Extraction for Systematic Reviews and Evidence Synthesis (v2.1.0)},
  year={2025},
  doi={10.5281/zenodo.18060972}
}"""

   
    try:
        citation_style = st.selectbox(
            "Select citation style",
            ["APA", "Harvard", "MLA", "Chicago", "IEEE", "Vancouver"]
        )
    except st.errors.StreamlitDuplicateElementId as e:
        st.warning(f"Citation section skipped (Duplicate ID): {e}")
        
        return

    if citation_style == "APA":
        citation_text = apa_citation
    elif citation_style == "Harvard":
        citation_text = harvard_citation
    elif citation_style == "MLA":
        citation_text = mla_citation
    elif citation_style == "Chicago":
        citation_text = chicago_citation
    elif citation_style == "IEEE":
        citation_text = ieee_citation
    elif citation_style == "Vancouver":
        citation_text = vancouver_citation

    escaped_citation = html.escape(citation_text)
    
    st.markdown(f'<div class="citation-box"><p style="margin:0; color: #F0F4F8;">{escaped_citation}</p></div>', unsafe_allow_html=True)

    js_citation_text = json.dumps(citation_text)
    
    st.markdown(f"""
    <div style="display:flex; gap:10px; margin-top:10px; margin-bottom:10px; position:relative; flex-wrap:wrap;" id="button-container">
        <button id="copy-btn" class="custom-button">Copy</button>
        <a download="ReviewAid_citation.ris" href="data:application/x-research-info-systems;base64,{base64.b64encode(ris_data.encode()).decode()}" class="custom-button">RIS Format</a>
        <a download="ReviewAid_citation.bib" href="data:application/x-bibtex;base64,{base64.b64encode(bib_data.encode()).decode()}" class="custom-button">BibTeX Format</a>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <script>
    function copyCitation() {{
        const citationText = {js_citation_text};
        navigator.clipboard.writeText(citationText).then(() => {{
            const btn = document.getElementById('copy-btn');
            const originalText = btn.innerText;
            btn.innerText = "Copied!";
            setTimeout(() => {{ 
                btn.innerText = originalText; 
            }}, 2000);
        }}).catch(err => {{
            console.error('Failed to copy text: ', err);
        }});
    }}
    
    document.addEventListener('DOMContentLoaded', function() {{
        const copyBtn = document.getElementById('copy-btn');
        if (copyBtn) {{
            copyBtn.addEventListener('click', copyCitation);
        }}
    }});
    </script>
    """, unsafe_allow_html=True)



def init_analytics():
    if db is None: return
    if "visit_recorded" not in st.session_state:
        doc_ref = db.collection("ReviewAidAnalytics").document("counters")
        doc_ref.update({"total_visits": firestore.Increment(1)})
        st.session_state.visit_recorded = True

def increment_firebase_counter(field):
    if db is None: return
    try:
        doc_ref = db.collection("ReviewAidAnalytics").document("counters")
        doc_ref.update({field: firestore.Increment(1)})
    except Exception as e: print(f"Analytics Error: {e}")

def update_processing_stats(mode, count=1):
    if db is None: return
    if mode == "screener":
        for _ in range(count): increment_firebase_counter("papers_screened")
    elif mode == "extractor":
        for _ in range(count): increment_firebase_counter("papers_extracted")

def update_terminal_log(msg, level="INFO"):
    if "terminal_logs" not in st.session_state or "terminal_placeholder" not in st.session_state:
        return

    timestamp = datetime.now().strftime("%H:%M:%S.%f")[:-3] 
    
    if level == "INFO":
        color_class = "log-info"
    elif level == "SUCCESS":
        color_class = "log-success"
    elif level == "WARN":
        color_class = "log-warn"
    elif level == "ERROR":
        color_class = "log-error"
    elif level == "SYSTEM":
        color_class = "log-system"
    elif level == "DEBUG":
        color_class = "log-debug"
    else:
        color_class = "log-info"
        
    log_entry = f'<span class="terminal-line"><span class="terminal-timestamp">[{timestamp}]</span><span class="{color_class}">{level}</span>: {html.escape(msg)}</span>'
    
    st.session_state.terminal_logs.append(log_entry)
    
    if len(st.session_state.terminal_logs) > MAX_LOG_ENTRIES:
        st.session_state.terminal_logs.pop(0)
    
    current_time = time.time()
    if current_time - st.session_state.last_log_update_time > 0.5:
        full_log_html = '<div class="terminal-container" id="terminal-container">' + "".join(st.session_state.terminal_logs) + '</div>'
        scroll_script = """
        <script>
            var element = document.getElementById("terminal-container");
            if(element){
                setTimeout(function() {
                    element.scrollTop = element.scrollHeight;
                }, 50);
            }
        </script>
        """
        try:
            st.session_state.terminal_placeholder.markdown(full_log_html + scroll_script, unsafe_allow_html=True)
        except Exception:
            pass
        st.session_state.last_log_update_time = current_time

def extract_pdf_content(pdf_bytes):
    try:
        update_terminal_log("Initializing PDF extraction engine (PyMuPDF)...", "DEBUG")
    except:
        pass
        
    doc = None
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = doc.page_count
        try:
            update_terminal_log(f"Document opened successfully. Total pages: {page_count}", "INFO")
        except:
            pass
        
        MAX_CHAR_LIMIT = 600000 
        full_text_parts = []
        current_length = 0
        references_found = False
        
        for i, page in enumerate(doc):
            if current_length > MAX_CHAR_LIMIT and not references_found:
                 try:
                     update_terminal_log(f"Token limit reached at Page {i+1}. Stopping PDF read.", "INFO")
                 except:
                     pass
                 break
            
            page_text = page.get_text()
            
            if re.search(r'(?:\n|\r\n){1,2}(References|Reference|Bibliography)(?:\s|\r?\n|$)', page_text, re.IGNORECASE):
                references_found = True 

            full_text_parts.append(page_text)
            current_length += len(page_text)
            
            del page_text 
            
        full_text = "".join(full_text_parts)
        del full_text_parts

        ref_match = re.search(r'(?:\n|\r\n){1,2}(References|Reference|Bibliography)(?:\s|\r?\n|$)', full_text, re.IGNORECASE)
        
        if ref_match:
            main_text_end = ref_match.start()
            full_text = full_text[:main_text_end]
            try:
                update_terminal_log("References section detected and removed to save tokens.", "INFO")
            except:
                pass
            del ref_match

        metadata = doc.metadata
        title = metadata.get('title', '')
        author = metadata.get('author', '')
        del metadata 
        
        try:
            update_terminal_log(f"Metadata read -> Title: '{title}', Author: '{author}'", "DEBUG")
        except:
            pass
    
        if not title or len(title) < 5 or "microsoft" in title.lower() or ".doc" in title.lower() or title.isdigit():
            try:
                update_terminal_log("Metadata title appears empty or invalid. Scanning first page text for title...", "DEBUG")
            except:
                pass
            
            first_page_text = full_text[:3000]
            lines = first_page_text.split('\n')
            
            for line in lines:
                clean_line = line.strip()

                if len(clean_line) > 10 and len(clean_line) < 200:
                    if not re.match(r'^[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}$', clean_line, re.IGNORECASE) and \
                       not clean_line.startswith("http") and \
                       "abstract" not in clean_line.lower() and \
                       "introduction" not in clean_line.lower() and \
                       "keywords" not in clean_line.lower() and \
                       "page" not in clean_line.lower() and \
                       "vol." not in clean_line.lower():
                        

                        title = clean_line
                        try:
                            update_terminal_log(f"Found candidate title in text: '{title[:50]}...'", "INFO")
                        except:
                            pass
                        break
        
        year = ''
        if not year:
            if page_count > 0:
                year_match = re.search(r'\b(19|20)\d{2}\b', full_text[:5000]) 
                if year_match:
                    year = year_match.group()
                    try:
                        update_terminal_log(f"Year found on Page 1: {year}", "DEBUG")
                    except:
                        pass
        
        if not year:
            creation_date = doc.metadata.get('creationDate', '')
            if creation_date:
                year_match = re.search(r'\b(19|20)\d{2}\b', creation_date)
                if year_match:
                    year = year_match.group()
                    try:
                        update_terminal_log(f"Year derived from creationDate: {year}", "DEBUG")
                    except:
                        pass
        
        return full_text, title, author, year
        
    except Exception as e:
        try:
            update_terminal_log(f"Error during PDF extraction: {str(e)}", "ERROR")
        except:
            pass
        return "", "", "", ""
    finally:
        if doc:
            doc.close()
            del doc

def preprocess_text_for_ai(text, max_tokens=MAX_INPUT_TOKENS_SCREENER):
    if "  " in text or "\n" in text:
        text = " ".join(text.split())
    
    char_limit = max_tokens * 4
    if len(text) > char_limit:
        try:
            update_terminal_log(f"Text exceeds token limit ({len(text)} > {char_limit}). Truncating...", "WARN")
        except:
            pass
        text = text[:char_limit]
    
    return text

def to_docx(df):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5) 
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    section = doc.sections[0]
    
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    logo_path = os.path.join("assets", "RA_transparent.png")
    if os.path.exists(logo_path):
        try:
            run_logo = header_para.add_run()
            run_logo.add_picture(logo_path, height=Inches(0.4))
        except Exception:
            pass 

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    generated_date = datetime.now().strftime("%B %d, %Y")
    
    footer_text = f"Generated by ReviewAid on {generated_date}"
    
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = 'Arial'
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    title = doc.add_heading('Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    num_cols = len(df.columns)
    table = doc.add_table(rows=1, cols=num_cols)

    table.style = 'Table Grid' 
    
    table.allow_autofit = True
    for row in table.rows:
        row.allow_break_across_pages = True

    usable_width = section.page_width - section.left_margin - section.right_margin - Inches(0.2)
    col_width = usable_width / num_cols if num_cols > 0 else usable_width

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        cell = hdr_cells[i]
        cell.text = str(col)
        cell.width = col_width 
        
        cell_par = cell.paragraphs[0]
        cell_par.runs[0].font.bold = True
        cell_par.runs[0].font.name = 'Arial'
        cell_par.runs[0].font.size = Pt(10)
        cell_par.runs[0].font.color.rgb = RGBColor(0, 0, 0) 
        cell_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FFFFFF') 
        cell._element.get_or_add_tcPr().append(shading_elm)
    
        tc_borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000') 
        tc_borders.append(bottom)
        cell._element.get_or_add_tcPr().append(tc_borders)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = row_cells[i]
            cell.width = col_width 
            
            para = cell.paragraphs[0]
            
            safe_val = str(val)
            safe_val = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', safe_val)
            
            run = para.add_run(safe_val)
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0) 
    
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'FFFFFF')
            tcPr.append(shading_elm)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def to_pdf(df):
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    
    pdf.set_auto_page_break(auto=True, margin=15) 
    
    pdf.set_font("Arial", size=8)

    logo_path = os.path.join("assets", "RA_transparent.png")
    if os.path.exists(logo_path):
        try:
            pdf.image(logo_path, x=250, y=10, w=30)
        except:
            pass 

    generated_date = datetime.now().strftime("%Y-%m-%d")

    col_width = 277 / len(df.columns) 
    line_height = 5 

    pdf.set_fill_color(65, 137, 220) 
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Arial", style='B', size=8)
    
    if pdf.get_y() > 260: 
        pdf.add_page()
        
    for col in df.columns:
        pdf.cell(col_width, 7, str(col), border=1, align='C', fill=True)
    pdf.ln(7)
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Arial", size=7)
    is_odd = True
    
    table_x_start = pdf.get_x()
    
    for index, row in df.iterrows():

        if pdf.get_y() > 250:
            pdf.add_page()

            pdf.set_y(10)
            pdf.set_fill_color(65, 137, 220) 
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Arial", style='B', size=8)
            for col in df.columns:
                pdf.cell(col_width, 7, str(col), border=1, align='C', fill=True)
            pdf.ln(7)
    
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", size=7)
            is_odd = True
            table_x_start = pdf.get_x()

        if not is_odd:
             pdf.set_fill_color(245, 245, 245)
        else:
             pdf.set_fill_color(255, 255, 255)

        y_start = pdf.get_y()
        max_y_end = y_start
        
        col_heights = []
        for col_name, val in zip(df.columns, row):
            clean_text = str(val).encode('latin-1', 'replace').decode('latin-1')
       
            chars_per_line = int(col_width * 3.5) 
            lines = len(clean_text) // chars_per_line + 1
            
            lines += clean_text.count('\n')
            
            cell_h = lines * line_height
            col_heights.append(cell_h)
            if cell_h > (max_y_end - y_start):
                max_y_end = y_start + cell_h
        

        row_height = max_y_end - y_start
        if row_height < line_height:
            row_height = line_height

        x_cursor = table_x_start
        
        for i, val in enumerate(row):
            clean_text = str(val).encode('latin-1', 'replace').decode('latin-1')

            pdf.set_xy(x_cursor, y_start)
            
            pdf.multi_cell(col_width, row_height, clean_text, border=1, fill=(not is_odd), align='L')
            
            x_cursor += col_width

        pdf.set_y(max_y_end)
        pdf.set_x(table_x_start) 
        
        is_odd = not is_odd

    pdf.set_y(-15)
    pdf.set_font("Arial", size=7, style='I')
    pdf.set_text_color(128, 128, 128)
    pdf.set_x(10) 
    pdf.cell(0, 10, f"Generated by ReviewAid on {generated_date}", 0, 0, 'L')
    
    return pdf.output(dest='S').encode('latin1')

def to_csv(df):
    buffer = io.StringIO()
    buffer.write("# Generated by ReviewAid\n")
    df.to_csv(buffer, index=False)
    return buffer.getvalue().encode('utf-8')

def to_excel(df):
    import xlsxwriter
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='ReviewAid Data')
        
        workbook = writer.book
        worksheet = writer.sheets['ReviewAid Data']
        
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#4189DC', 
            'font_color': '#FFFFFF',
            'border': 1
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            
        generated_date = datetime.now().strftime("%Y-%m-%d")

        worksheet.set_footer(f'&LGenerated by ReviewAid on {generated_date}')
        
    buffer.seek(0)
    return buffer.getvalue()



class BaseLLMProvider(ABC):
    """Abstract base class for LLM providers."""
    def __init__(self, api_key: str, model_name: str, **kwargs):
        self.api_key = api_key
        self.model_name = model_name
        self.extra_params = kwargs

    @abstractmethod
    def generate(self, messages: List[Dict[str, str]], temperature: float, max_tokens: int) -> str:
        pass

class OpenAIProvider(BaseLLMProvider):
    def __init__(self, api_key, model_name="gpt-4o", base_url=None):
        super().__init__(api_key, model_name)
        if not OpenAI: raise ImportError("OpenAI library not installed. Run: pip install openai")
        self.client = OpenAI(api_key=api_key, base_url=base_url)

    def generate(self, messages, temperature, max_tokens):
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content

class AnthropicProvider(BaseLLMProvider):
    def __init__(self, api_key, model_name="claude-sonnet-4-20250514"):
        super().__init__(api_key, model_name)
        if not Anthropic: raise ImportError("Anthropic library not installed. Run: pip install anthropic")
        self.client = Anthropic(api_key=api_key)

    def generate(self, messages, temperature, max_tokens):
        system_content = ""
        user_messages = []
        
        for msg in messages:
            if msg['role'] == 'system':
                system_content = msg['content']
            else:
                user_messages.append(msg)
        
        response = self.client.messages.create(
            model=self.model_name,
            system=system_content,
            messages=user_messages,
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.content[0].text

class CohereProvider(BaseLLMProvider):
    def __init__(self, api_key, model_name="command-a-03-2025"):
        super().__init__(api_key, model_name)
        if not cohere: raise ImportError("Cohere library not installed. Run: pip install cohere")
        self.client = cohere.Client(api_key=api_key)

    def generate(self, messages, temperature, max_tokens):
        chat_history = []
        message_content = ""
        
        for msg in messages:
            if msg['role'] == 'user':
                if not message_content: message_content = msg['content']
            elif msg['role'] == 'assistant':
                chat_history.append({"role": "CHATBOT", "message": msg['content']})
            elif msg['role'] == 'system':
                message_content = msg['content'] + "\n\n" + message_content

        response = self.client.chat(
            message=message_content,
            chat_history=chat_history,
            temperature=temperature,
            max_tokens=max_tokens,
            model=self.model_name
        )
        return response.text

class DeepSeekProvider(BaseLLMProvider):
    """DeepSeek is largely OpenAI compatible"""
    def __init__(self, api_key, model_name="deepseek-chat"):
        super().__init__(api_key, model_name)
        if not OpenAI: raise ImportError("OpenAI library required for DeepSeek. Run: pip install openai")
        self.client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

    def generate(self, messages, temperature, max_tokens):
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content

class GLMProvider(BaseLLMProvider):
    """Z.ai (GLM) provider using existing ZaiClient logic"""
    def __init__(self, api_key, model_name="GLM-4.6V-Flash"):
        super().__init__(api_key, model_name)
        if not ZaiClient: raise ImportError("Zai library not found.")
        self.client = ZaiClient(api_key=api_key)

    def generate(self, messages, temperature, max_tokens):
        prompt = ""
        for msg in messages:
            role = msg['role'].upper()
            content = msg['content']
            prompt += f"[{role}]: {content}\n"
            
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content

class OllamaProvider(BaseLLMProvider):
    def __init__(self, api_key, model_name="llama3", base_url="http://localhost:11434"):
        super().__init__(api_key, model_name, base_url=base_url)
        if not ollama: raise ImportError("Ollama library not installed. Run: pip install ollama")
        self.client = ollama.Client(host=base_url)

    def generate(self, messages, temperature, max_tokens):
        response = self.client.chat(
            model=self.model_name,
            messages=messages,
            options={
                'temperature': temperature,
                'num_predict': max_tokens
            }
        )
        return response['message']['content']

def get_provider_instance(provider_name: str, api_key: str, model_name: str, **kwargs):
    """Factory function to get the provider instance."""
    provider_map = {
        "OpenAI": OpenAIProvider,
        "Anthropic": AnthropicProvider,
        "Cohere": CohereProvider,
        "DeepSeek": DeepSeekProvider,
        "GLM (Z.ai)": GLMProvider,
        "Ollama (Local)": OllamaProvider
    }
    
    provider_class = provider_map.get(provider_name)
    if not provider_class:
        raise ValueError(f"Provider {provider_name} not supported.")
    
    return provider_class(api_key=api_key, model_name=model_name, **kwargs)

def query_llm(prompt, provider_name, api_key, model_name, temperature=0.1, max_tokens=8192):
    """
    Generic query function replacing query_zai.
    Handles retries and delegates to specific providers.
    """
    if not api_key and provider_name != "Ollama (Local)":
        return None
    
    extra_args = {}
    if provider_name == "Ollama (Local)":
        extra_args['base_url'] = st.session_state.get('ollama_base_url', 'http://localhost:11434')

    try:
        update_terminal_log(f"Initializing {provider_name} Client...", "DEBUG")
    except: pass

    max_retries = 10 
    provider = None

    for attempt in range(max_retries):
        try:
            try:
                update_terminal_log(f"LLM Call Attempt {attempt + 1}/{max_retries} via {provider_name}...", "DEBUG")
            except: pass
            
          
            provider = get_provider_instance(provider_name, api_key, model_name, **extra_args)
            
            messages = [{"role": "user", "content": prompt}]
            
            result_content = provider.generate(messages, temperature, max_tokens)
            
            del provider 
            try: update_terminal_log("Response received successfully.", "SUCCESS")
            except: pass
            return result_content
            
        except Exception as e:
            error_str = str(e)
            is_rate_limit = False
            
        
            rate_keywords = ["429", "rate limit", "too many requests", "quota", "overload", "rate_limit_exceeded"]
            if any(k in error_str.lower() for k in rate_keywords):
                is_rate_limit = True
            
            if is_rate_limit:
                try: update_terminal_log(f"Rate Limit / Quota Exceeded detected.", "WARN")
                except: pass
                
                wait_time = 15 * (2 ** attempt)
                if attempt < max_retries - 1:
                    try: update_terminal_log(f"Rate limit hit. Waiting {wait_time} seconds before retry...", "WARN")
                    except: pass
                    time.sleep(wait_time)
                    try: update_terminal_log(f"Resuming retry...", "INFO")
                    except: pass
                else:
                    try: update_terminal_log("Max retries reached for rate limit.", "ERROR")
                    except: pass
                    return None
            else:
                try: update_terminal_log(f"API Error: {error_str}", "ERROR")
                except: pass
                if attempt < 3: time.sleep(2)
                else: return None
    
    return None